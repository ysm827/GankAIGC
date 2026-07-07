import json
import socket
import sys
import types
import zipfile
from io import BytesIO

import httpx

import app.config as config_module
from app.database import SessionLocal
from app.models.models import CreditTransaction, OptimizationSegment, OptimizationSession, User
from app.utils.auth import create_user_access_token, get_password_hash, verify_stream_token


class NoRunBackgroundTasks:
    def add_task(self, *args, **kwargs):
        return None


async def _noop_run_optimization(*args, **kwargs):
    return None


class ReadyZhuqueService:
    async def readiness(self, text=None):
        return {
            "ready": True,
            "connected": True,
            "page_found": True,
            "has_token": True,
            "remaining_uses": 20,
            "button_enabled": True,
            "text_length": len(text or ""),
            "text_length_ok": True,
            "estimated_first_round_credits": 10,
            "estimated_max_round_credits": 50,
            "message": "朱雀已就绪",
            "actions": [],
        }


def _allow_public_model_url_dns(monkeypatch):
    monkeypatch.setattr(
        socket,
        "getaddrinfo",
        lambda *args, **kwargs: [(socket.AF_INET, socket.SOCK_STREAM, 6, "", ("8.8.8.8", 443))],
    )


def _create_user(credit_balance=0, is_unlimited=False):
    db = SessionLocal()
    try:
        user = User(
            username="alice",
            password_hash=get_password_hash("Password123!"),
            access_link="http://testserver/access/alice",
            is_active=True,
            credit_balance=credit_balance,
            is_unlimited=is_unlimited,
        )
        db.add(user)
        db.commit()
        db.refresh(user)
        token = create_user_access_token(user.id, user.username)
        return user.id, token
    finally:
        db.close()


def test_calculate_optimization_credits_uses_billable_characters_and_stage_multiplier():
    from app.services.credit_service import calculate_optimization_credits, count_billable_characters

    assert count_billable_characters("a b\nc") == 3
    assert calculate_optimization_credits("短文本", "paper_enhance") == 1
    assert calculate_optimization_credits("汉" * 1000, "paper_enhance") == 1
    assert calculate_optimization_credits("汉" * 1001, "paper_enhance") == 2
    assert calculate_optimization_credits("汉" * 3000, "paper_polish_enhance") == 6


def test_parse_markdown_document_upload_returns_editable_text(client):
    _, token = _create_user(credit_balance=0)

    response = client.post(
        "/api/optimization/documents/parse",
        files={"file": ("paper.md", b"# Title\n\nMarkdown body", "text/markdown")},
        headers={"Authorization": f"Bearer {token}"},
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["filename"] == "paper.md"
    assert payload["parser"] == "markdown"
    assert payload["text"] == "# Title\n\nMarkdown body"
    assert payload["char_count"] > 0


def test_parse_txt_document_upload_returns_plain_text(client):
    _, token = _create_user(credit_balance=0)

    response = client.post(
        "/api/optimization/documents/parse",
        files={"file": ("paper.txt", "标题\n\n正文段落".encode("gb18030"), "text/plain")},
        headers={"Authorization": f"Bearer {token}"},
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["filename"] == "paper.txt"
    assert payload["parser"] == "plain_text"
    assert payload["document_format"] == "txt"
    assert payload["parse_engine"] == "plain_text"
    assert payload["text"] == "标题\n\n正文段落"
    assert payload["char_count"] > 0


def test_parse_document_upload_zero_config_falls_back_to_default_limit(client, monkeypatch):
    from app.routes import optimization

    _, token = _create_user(credit_balance=0)
    monkeypatch.setattr(optimization.settings, "MAX_UPLOAD_FILE_SIZE_MB", 0, raising=False)

    response = client.post(
        "/api/optimization/documents/parse",
        files={"file": ("paper.md", b"# Title\n\nMarkdown body", "text/markdown")},
        headers={"Authorization": f"Bearer {token}"},
    )

    assert response.status_code == 200
    assert response.json()["text"] == "# Title\n\nMarkdown body"


def test_parse_docx_document_upload_uses_docx_styles_before_text_rules(client):
    from docx import Document

    _, token = _create_user(credit_balance=0)
    document = Document()
    document.add_paragraph("论文标题", style="Title")
    document.add_paragraph("摘要", style="Heading 1")
    document.add_paragraph("摘要正文用于说明研究背景和主要结论。" * 4)
    document.add_paragraph("1 绪论", style="Heading 1")
    document.add_paragraph("正文段落围绕研究背景展开，且长度足够进入正文候选。" * 5)
    buffer = BytesIO()
    document.save(buffer)

    response = client.post(
        "/api/optimization/documents/parse",
        files={
            "file": (
                "paper.docx",
                buffer.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        headers={"Authorization": f"Bearer {token}"},
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["filename"] == "paper.docx"
    assert payload["parser"] == "python_docx"
    assert payload["parse_engine"] == "python_docx"
    assert payload["structure_summary"]["TITLE"] == 1
    assert payload["structure_summary"]["ABSTRACT_HEADING"] == 1
    assert payload["structure_summary"]["ABSTRACT_BODY"] == 1
    assert payload["structure_summary"]["SECTION_HEADING"] == 1
    body_segments = [segment for segment in payload["segments"] if segment["semantic_type"] == "BODY"]
    assert len(body_segments) == 1
    assert body_segments[0]["reduce_allowed"] is True


def test_parse_docx_document_upload_falls_back_to_markitdown(client, monkeypatch):
    from app.services import document_structure_service as structure_module

    _, token = _create_user(credit_balance=0)
    seen_extensions = []

    class FakeMarkItDown:
        def __init__(self, enable_plugins=False):
            assert enable_plugins is False

        def convert_stream(self, stream, *, file_extension=None):
            seen_extensions.append(file_extension)
            assert stream.read().startswith(b"PK\x03\x04")
            return types.SimpleNamespace(text_content="# 论文标题\n\n解析正文")

    monkeypatch.setitem(sys.modules, "markitdown", types.SimpleNamespace(MarkItDown=FakeMarkItDown))
    monkeypatch.setattr(structure_module, "classify_docx_style", lambda _style_name: (_ for _ in ()).throw(RuntimeError("docx broken")))

    response = client.post(
        "/api/optimization/documents/parse",
        files={
            "file": (
                "paper.docx",
                b"PK\x03\x04 fake docx content",
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        headers={"Authorization": f"Bearer {token}"},
    )

    assert response.status_code == 200
    payload = response.json()
    assert seen_extensions == [".docx"]
    assert payload["filename"] == "paper.docx"
    assert payload["parser"] == "markitdown"
    assert payload["parse_engine"] == "markitdown"
    assert payload["text"] == "# 论文标题\n\n解析正文"
    assert payload["char_count"] == 8


def test_parse_pdf_document_upload_uses_mineru_content_list(client, monkeypatch):
    from app.services import document_structure_service as structure_module
    from app.services import mineru_service as mineru_module

    _, token = _create_user(credit_balance=0)
    seen_calls = []

    def fake_parse_pdf(filename, content):
        seen_calls.append((filename, content[:4]))
        return mineru_module.MinerUParseResult(
            batch_id="batch-123",
            trace_id="trace-abc",
            model_version="vlm",
            content_list_name="paper_content_list.json",
            full_zip_url="https://mineru.example/full.zip",
            content_items=[
                {"type": "text", "text": "1 绪论", "text_level": 1, "bbox": [1, 2, 3, 4], "page_idx": 0},
                {
                    "type": "text",
                    "text": "正文段落围绕研究背景展开，且长度足够进入正文候选。" * 5,
                    "bbox": [1, 2, 3, 4],
                    "page_idx": 0,
                },
                {
                    "type": "table",
                    "table_caption": ["表 1 实验结果"],
                    "table_body": "<table><tr><td>A</td></tr></table>",
                    "table_footnote": [],
                    "bbox": [1, 2, 3, 4],
                    "page_idx": 1,
                },
                {"type": "equation", "text": "$$x=y$$", "bbox": [1, 2, 3, 4], "page_idx": 1},
                {
                    "type": "list",
                    "sub_type": "ref_text",
                    "list_items": ["[1] Author. Title[J]. 2020."],
                    "bbox": [1, 2, 3, 4],
                    "page_idx": 2,
                },
            ],
        )

    monkeypatch.setattr(structure_module.settings, "PDF_STRUCTURE_ENGINE", "mineru", raising=False)
    monkeypatch.setattr(structure_module.mineru_service, "parse_pdf", fake_parse_pdf)

    response = client.post(
        "/api/optimization/documents/parse",
        files={"file": ("paper.pdf", b"%PDF-1.7 fake pdf", "application/pdf")},
        headers={"Authorization": f"Bearer {token}"},
    )

    assert response.status_code == 200
    payload = response.json()
    assert seen_calls == [("paper.pdf", b"%PDF")]
    assert payload["filename"] == "paper.pdf"
    assert payload["parser"] == "mineru"
    assert payload["parse_engine"] == "mineru"
    assert payload["parse_fallback_used"] is False
    trace = json.loads(payload["parse_trace"])
    assert trace["engine"] == "mineru"
    assert trace["api"] == "v4"
    assert trace["batch_id"] == "batch-123"
    assert trace["content_item_count"] == 5
    assert trace["content_list_name"] == "paper_content_list.json"
    assert payload["structure_summary"]["SECTION_HEADING"] == 1
    assert payload["structure_summary"]["BODY"] == 1
    assert payload["structure_summary"]["TABLE"] == 1
    assert payload["structure_summary"]["FORMULA"] == 1
    assert payload["structure_summary"]["REFERENCE_ITEM"] == 1
    by_type = {segment["semantic_type"]: segment for segment in payload["segments"]}
    assert by_type["BODY"]["reduce_allowed"] is True
    assert by_type["SECTION_HEADING"]["reduce_allowed"] is False
    assert by_type["TABLE"]["reduce_allowed"] is False
    assert by_type["FORMULA"]["reduce_allowed"] is False
    assert by_type["REFERENCE_ITEM"]["reduce_allowed"] is False
    assert by_type["BODY"]["page_number"] == 1
    assert by_type["TABLE"]["page_number"] == 2
    assert by_type["REFERENCE_ITEM"]["page_number"] == 3


def test_mineru_pdf_cleanup_removes_repeated_headers_without_losing_references():
    from app.services import document_structure_service as structure_module
    from app.services import mineru_service as mineru_module

    result = mineru_module.MinerUParseResult(
        batch_id="batch-refs",
        trace_id="trace-refs",
        model_version="vlm",
        content_list_name="journal_content_list.json",
        full_zip_url="https://mineru.example/full.zip",
        content_items=[
            {"type": "text", "text": "8 References", "text_level": 1, "page_idx": 30},
            {
                "type": "text",
                "text": "[1] Chen, A. (2021). Cultural communication in digital platforms. Journal of Media Studies, 12(1), 1-10.",
                "page_idx": 30,
            },
            {"type": "text", "text": "32", "page_idx": 31},
            {"type": "text", "text": "International Journal of Social Science and Education Research", "page_idx": 31},
            {"type": "text", "text": "Volume 7 Issue 9, 2024", "page_idx": 31},
            {"type": "text", "text": "ISSN: 2637-6067", "page_idx": 31},
            {"type": "text", "text": "DOI: 10.6918/IJOSSER.202409_7(9).0039", "page_idx": 31},
            {
                "type": "text",
                "text": "[2] Zhang, L. (2020). Digital culture and media practice. Journal of Communication, 8(2), 20-30. doi:10.1000/ref.2020",
                "page_idx": 31,
            },
            {"type": "text", "text": "International Journal of Social Science and Education Research", "page_idx": 32},
            {"type": "text", "text": "Volume 7 Issue 9, 2024", "page_idx": 32},
            {"type": "text", "text": "ISSN: 2637-6067", "page_idx": 32},
            {"type": "text", "text": "DOI: 10.6918/IJOSSER.202409_7(9).0039", "page_idx": 32},
        ],
    )

    parsed = structure_module.build_parsed_document_from_mineru_result(result)

    assert "International Journal of Social Science and Education Research" not in parsed.text
    assert "Volume 7 Issue 9, 2024" not in parsed.text
    assert "ISSN: 2637-6067" not in parsed.text
    assert "10.6918/IJOSSER.202409_7(9).0039" not in parsed.text
    assert "\n\n32\n\n" not in f"\n\n{parsed.text}\n\n"
    assert "[1] Chen, A. (2021)" in parsed.text
    assert "[2] Zhang, L. (2020)" in parsed.text

    semantic_by_text = {segment.text: segment.semantic_type for segment in parsed.segments}
    assert semantic_by_text["8 References"] == structure_module.SEMANTIC_TYPE_REFERENCE_HEADING
    assert semantic_by_text[
        "[1] Chen, A. (2021). Cultural communication in digital platforms. Journal of Media Studies, 12(1), 1-10."
    ] == structure_module.SEMANTIC_TYPE_REFERENCE_ITEM
    assert semantic_by_text[
        "[2] Zhang, L. (2020). Digital culture and media practice. Journal of Communication, 8(2), 20-30. doi:10.1000/ref.2020"
    ] == structure_module.SEMANTIC_TYPE_REFERENCE_ITEM
    assert all(segment.semantic_type != structure_module.SEMANTIC_TYPE_HEADER_FOOTER for segment in parsed.segments)


def test_mineru_pdf_maps_real_ref_text_items_from_content_list():
    from app.services import document_structure_service as structure_module
    from app.services import mineru_service as mineru_module

    result = mineru_module.MinerUParseResult(
        batch_id="batch-real-ref-text",
        trace_id="trace-real-ref-text",
        model_version="vlm",
        content_list_name="real_content_list.json",
        full_zip_url="https://mineru.example/full.zip",
        content_items=[
            {"type": "text", "text": "References", "text_level": 2, "page_idx": 4},
            {
                "type": "ref_text",
                "text": "[1] Sijing, C., & Lie, Z. (2020, December). Research on Cultural and Creative Products Combined with Mobile Virtual Applications for Chinese Museums. In 2020 International Conference on Intelligent Design (ICID) (pp. 267-272). IEEE.",
                "page_idx": 4,
            },
            {"type": "header", "text": "International Journal of Social Science and Education Research", "page_idx": 4},
            {"type": "page_number", "text": "32", "page_idx": 4},
            {
                "type": "ref_text",
                "text": "[2] Zhang, L. (2020). Foreshadowing the Future of Capitalism: Surveillance Technology and Digital Realism in Xu Bing’s Dragonfly Eyes. Comparative Cinema, 8(14), 62-81.",
                "page_idx": 5,
            },
        ],
    )

    parsed = structure_module.build_parsed_document_from_mineru_result(result)

    assert "[1] Sijing, C." in parsed.text
    assert "[2] Zhang, L." in parsed.text
    assert "International Journal of Social Science and Education Research" not in parsed.text
    assert "\n\n32\n\n" not in f"\n\n{parsed.text}\n\n"

    ref_segments = [segment for segment in parsed.segments if segment.semantic_type == structure_module.SEMANTIC_TYPE_REFERENCE_ITEM]
    assert [segment.page_number for segment in ref_segments] == [5, 6]
    assert all(segment.reduce_allowed is False for segment in ref_segments)
    assert all(segment.semantic_reason == "mineru_reference_text" for segment in ref_segments)


def test_parse_pdf_document_upload_falls_back_to_markitdown_when_mineru_fails(client, monkeypatch):
    from app.services import document_structure_service as structure_module
    from app.services import mineru_service as mineru_module

    _, token = _create_user(credit_balance=0)
    seen_extensions = []

    class FakeMarkItDown:
        def __init__(self, enable_plugins=False):
            assert enable_plugins is False

        def convert_stream(self, stream, *, file_extension=None):
            seen_extensions.append(file_extension)
            assert stream.read().startswith(b"%PDF")
            return types.SimpleNamespace(text_content="PDF fallback body")

    monkeypatch.setattr(structure_module.settings, "PDF_STRUCTURE_ENGINE", "mineru", raising=False)
    monkeypatch.setattr(
        structure_module.mineru_service,
        "parse_pdf",
        lambda _filename, _content: (_ for _ in ()).throw(mineru_module.MinerUConfigError("未配置 MINERU_API_TOKEN")),
    )
    monkeypatch.setitem(sys.modules, "markitdown", types.SimpleNamespace(MarkItDown=FakeMarkItDown))

    response = client.post(
        "/api/optimization/documents/parse",
        files={"file": ("paper.pdf", b"%PDF-1.7 fake pdf", "application/pdf")},
        headers={"Authorization": f"Bearer {token}"},
    )

    assert response.status_code == 200
    payload = response.json()
    assert seen_extensions == [".pdf"]
    assert payload["filename"] == "paper.pdf"
    assert payload["parser"] == "markitdown"
    assert payload["parse_engine"] == "markitdown"
    assert payload["parse_fallback_used"] is True
    assert "MinerU" in payload["warnings"][0]
    trace = json.loads(payload["parse_trace"])
    assert trace["engine"] == "markitdown"
    assert trace["fallback_from"] == "mineru"
    assert trace["fallback_reason"] == "MinerUConfigError"
    assert payload["text"] == "PDF fallback body"


def test_mineru_service_uses_v4_upload_poll_and_content_list_zip(monkeypatch):
    from app.services import mineru_service as mineru_module

    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, "w") as archive:
        archive.writestr("paper_content_list.json", json.dumps([
            {"type": "text", "text": "正文段落", "page_idx": 0, "bbox": [1, 2, 3, 4]}
        ]))
    zip_bytes = zip_buffer.getvalue()
    calls = []

    def fake_response(status_code, *, json_payload=None, content=None, url="https://mineru.test/fake"):
        return httpx.Response(
            status_code,
            json=json_payload,
            content=content,
            request=httpx.Request("GET", url),
        )

    class FakeClient:
        def __init__(self, timeout=None):
            self.timeout = timeout

        def __enter__(self):
            return self

        def __exit__(self, exc_type, exc, tb):
            return False

        def post(self, url, *, headers=None, json=None):
            calls.append(("POST", url, headers, json))
            return fake_response(
                200,
                json_payload={
                    "code": 0,
                    "data": {
                        "batch_id": "batch-1",
                        "file_urls": ["https://upload.example/paper"],
                    },
                    "msg": "ok",
                        "trace_id": "trace-1",
                },
                url=url,
            )

        def put(self, url, *, content=None):
            calls.append(("PUT", url, content))
            return fake_response(200, url=url)

        def get(self, url, *, headers=None):
            calls.append(("GET", url, headers))
            if url.endswith("/api/v4/extract-results/batch/batch-1"):
                return fake_response(
                    200,
                    json_payload={
                        "code": 0,
                        "data": {
                            "batch_id": "batch-1",
                            "extract_result": {
                                "file_name": "paper.pdf",
                                "state": "done",
                                "full_zip_url": "https://download.example/full.zip",
                                "err_msg": "",
                                "data_id": "",
                            },
                        },
                        "msg": "ok",
                    },
                    url=url,
                )
            return fake_response(200, content=zip_bytes, url=url)

    monkeypatch.setattr(mineru_module.settings, "MINERU_API_TOKEN", "token-1", raising=False)
    monkeypatch.setattr(mineru_module.settings, "MINERU_BASE_URL", "https://mineru.test", raising=False)
    monkeypatch.setattr(mineru_module.settings, "MINERU_MODEL_VERSION", "vlm", raising=False)
    monkeypatch.setattr(mineru_module.settings, "MINERU_LANGUAGE", "ch", raising=False)
    monkeypatch.setattr(mineru_module.settings, "MINERU_IS_OCR", False, raising=False)
    monkeypatch.setattr(mineru_module.settings, "MINERU_ENABLE_FORMULA", True, raising=False)
    monkeypatch.setattr(mineru_module.settings, "MINERU_ENABLE_TABLE", True, raising=False)
    monkeypatch.setattr(mineru_module.httpx, "Client", FakeClient)

    result = mineru_module.MinerUService().parse_pdf("paper.pdf", b"%PDF-body")

    assert result.batch_id == "batch-1"
    assert result.trace_id == "trace-1"
    assert result.content_list_name == "paper_content_list.json"
    assert result.content_items == [{"type": "text", "text": "正文段落", "page_idx": 0, "bbox": [1, 2, 3, 4]}]
    post_call = calls[0]
    assert post_call[0:2] == ("POST", "https://mineru.test/api/v4/file-urls/batch")
    assert post_call[2]["Authorization"] == "Bearer token-1"
    assert post_call[3]["files"] == [{"name": "paper.pdf", "is_ocr": False}]
    assert post_call[3]["model_version"] == "vlm"
    assert post_call[3]["enable_formula"] is True
    assert post_call[3]["enable_table"] is True
    assert calls[1] == ("PUT", "https://upload.example/paper", b"%PDF-body")
    assert calls[2][0:2] == ("GET", "https://mineru.test/api/v4/extract-results/batch/batch-1")
    assert calls[3][0:2] == ("GET", "https://download.example/full.zip")


def test_mineru_service_accepts_batch_extract_result_array(monkeypatch):
    from app.services import mineru_service as mineru_module

    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, "w") as archive:
        archive.writestr("paper_content_list.json", json.dumps([
            {"type": "text", "text": "数组形态解析正文", "page_idx": 0, "bbox": [1, 2, 3, 4]}
        ]))
    zip_bytes = zip_buffer.getvalue()

    def fake_response(status_code, *, json_payload=None, content=None, url="https://mineru.test/fake"):
        return httpx.Response(
            status_code,
            json=json_payload,
            content=content,
            request=httpx.Request("GET", url),
        )

    class FakeClient:
        def __init__(self, timeout=None):
            self.timeout = timeout

        def __enter__(self):
            return self

        def __exit__(self, exc_type, exc, tb):
            return False

        def post(self, url, *, headers=None, json=None):
            return fake_response(
                200,
                json_payload={
                    "code": 0,
                    "data": {
                        "batch_id": "batch-array",
                        "file_urls": ["https://upload.example/paper"],
                    },
                    "msg": "ok",
                    "trace_id": "trace-array",
                },
                url=url,
            )

        def put(self, url, *, content=None):
            return fake_response(200, url=url)

        def get(self, url, *, headers=None):
            if url.endswith("/api/v4/extract-results/batch/batch-array"):
                return fake_response(
                    200,
                    json_payload={
                        "code": 0,
                        "data": {
                            "batch_id": "batch-array",
                            "extract_result": [
                                {
                                    "file_name": "paper.pdf",
                                    "state": "done",
                                    "full_zip_url": "https://download.example/full.zip",
                                    "err_msg": "",
                                    "data_id": "",
                                }
                            ],
                        },
                        "msg": "ok",
                    },
                    url=url,
                )
            return fake_response(200, content=zip_bytes, url=url)

    monkeypatch.setattr(mineru_module.settings, "MINERU_API_TOKEN", "token-1", raising=False)
    monkeypatch.setattr(mineru_module.settings, "MINERU_BASE_URL", "https://mineru.test", raising=False)
    monkeypatch.setattr(mineru_module.httpx, "Client", FakeClient)

    result = mineru_module.MinerUService().parse_pdf("paper.pdf", b"%PDF-body")

    assert result.batch_id == "batch-array"
    assert result.trace_id == "trace-array"
    assert result.full_zip_url == "https://download.example/full.zip"
    assert result.content_items == [{"type": "text", "text": "数组形态解析正文", "page_idx": 0, "bbox": [1, 2, 3, 4]}]


def test_parse_pdf_line_segments_merge_same_page_body_lines():
    from app.services import document_structure_service as structure_module

    raw_segments = [
        (
            "1 Introduction",
            structure_module.SegmentSemanticDecision(
                structure_module.SEMANTIC_TYPE_SECTION_HEADING,
                structure_module.SEMANTIC_SOURCE_TEXT_RULE,
                0.9,
                False,
                "section_heading",
            ),
            1,
            '{"l": 1}',
        ),
        ("To popularize Chinese culture in digital media, various aspects need to be considered. The", None, 1, '{"line": 1}'),
        ("influence of gender-switching videos in the Chinese digital media environment is growing,", None, 1, '{"line": 2}'),
        ("indicating the potential for increased popularity of pop or queer culture [1]. Additionally, the", None, 1, '{"line": 3}'),
        ("construction of digital identities by transnational bloggers on Chinese social media platforms", None, 1, '{"line": 4}'),
        ("2 Methods", structure_module.SegmentSemanticDecision(
            structure_module.SEMANTIC_TYPE_SECTION_HEADING,
            structure_module.SEMANTIC_SOURCE_TEXT_RULE,
            0.9,
            False,
            "section_heading",
        ), 1, '{"l": 2}'),
    ]

    merged = structure_module.merge_pdf_line_segments(raw_segments)
    parsed = structure_module.build_parsed_document_from_raw_segments(
        merged,
        parser="markitdown",
        document_format="pdf",
        fallback_source=structure_module.SEMANTIC_SOURCE_MARKITDOWN_TEXT_RULE,
        warnings=[],
        parse_engine="markitdown",
        trace={"engine": "markitdown"},
    )

    assert len(parsed.segments) == 3
    assert parsed.segments[1].semantic_type == "BODY"
    assert "The\n\ninfluence" not in parsed.text
    assert "The influence of gender-switching videos" in parsed.segments[1].text
    assert parsed.segments[1].bbox_json is None


def test_parse_pdf_text_export_merges_soft_wrapped_lines():
    from app.services import document_structure_service as structure_module

    text = (
        "1 Introduction\n"
        "To popularize Chinese culture in digital media, various aspects need to be considered. The\n"
        "influence of gender-switching videos in the Chinese digital media environment is growing,\n"
        "indicating the potential for increased popularity of pop or queer culture [1]. Additionally, the\n"
        "construction of digital identities by transnational bloggers on Chinese social media platforms\n\n"
        "2 Methods\n"
        "The study collects video samples and interview material to compare narrative strategies."
    )

    parsed = structure_module.build_parsed_document_from_text(
        structure_module.normalize_parsed_document_text(text),
        parser="markitdown",
        document_format="pdf",
        semantic_source=structure_module.SEMANTIC_SOURCE_MARKITDOWN_TEXT_RULE,
        warnings=[],
        parse_engine="markitdown",
        trace={"engine": "markitdown"},
    )

    assert "The\n\ninfluence" not in parsed.text
    assert "The influence of gender-switching videos" in parsed.text
    assert len(parsed.segments) == 4
    assert parsed.segments[1].semantic_type == "BODY"
    assert parsed.segments[1].reduce_allowed is True


def test_parse_pdf_document_upload_extracts_text_with_real_markitdown(client, monkeypatch):
    from app.services import document_structure_service as structure_module

    _, token = _create_user(credit_balance=0)
    monkeypatch.setattr(structure_module.settings, "PDF_STRUCTURE_ENGINE", "markitdown", raising=False)
    pdf = b"""%PDF-1.4
1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj
2 0 obj << /Type /Pages /Kids [3 0 R] /Count 1 >> endobj
3 0 obj << /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >> endobj
4 0 obj << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> endobj
5 0 obj << /Length 44 >> stream
BT /F1 24 Tf 72 720 Td (PDF parsed body) Tj ET
endstream endobj
xref
0 6
0000000000 65535 f
0000000009 00000 n
0000000058 00000 n
0000000115 00000 n
0000000274 00000 n
0000000344 00000 n
trailer << /Root 1 0 R /Size 6 >>
startxref
438
%%EOF
"""

    response = client.post(
        "/api/optimization/documents/parse",
        files={"file": ("paper.pdf", pdf, "application/pdf")},
        headers={"Authorization": f"Bearer {token}"},
    )

    assert response.status_code == 200
    payload = response.json()
    assert payload["parser"] == "markitdown"
    assert payload["text"] == "PDF parsed body"


def test_parse_document_upload_rejects_unsupported_file_type(client):
    _, token = _create_user(credit_balance=0)

    response = client.post(
        "/api/optimization/documents/parse",
        files={"file": ("paper.xlsx", b"xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")},
        headers={"Authorization": f"Bearer {token}"},
    )

    assert response.status_code == 400
    assert "PDF(.pdf)、Word(.docx)、Markdown(.md/.markdown) 和 TXT(.txt)" in response.json()["detail"]


def test_parse_doc_document_upload_rejects_with_conversion_hint(client):
    _, token = _create_user(credit_balance=0)

    response = client.post(
        "/api/optimization/documents/parse",
        files={"file": ("paper.doc", b"doc", "application/msword")},
        headers={"Authorization": f"Bearer {token}"},
    )

    assert response.status_code == 400
    assert response.json()["detail"] == "暂不支持老版 Word(.doc)，请另存为 .docx 后上传"


def test_platform_mode_holds_character_based_credits_before_processing(client, monkeypatch):
    from app.routes import optimization

    user_id, token = _create_user(credit_balance=8)
    monkeypatch.setattr(optimization, "BackgroundTasks", NoRunBackgroundTasks)
    monkeypatch.setattr(optimization, "run_optimization", _noop_run_optimization)

    response = client.post(
        "/api/optimization/start",
        json={
            "original_text": "汉" * 3000,
            "processing_mode": "paper_polish_enhance",
            "billing_mode": "platform",
        },
        headers={"Authorization": f"Bearer {token}"},
    )

    assert response.status_code == 200
    assert response.json()["charge_status"] == "held"
    assert response.json()["charged_credits"] == 6

    db = SessionLocal()
    try:
        user = db.query(User).filter(User.id == user_id).one()
        session = db.query(OptimizationSession).filter(OptimizationSession.user_id == user_id).one()
        transaction = db.query(CreditTransaction).filter(CreditTransaction.user_id == user_id).one()
        assert user.credit_balance == 2
        assert session.billing_mode == "platform"
        assert transaction.delta == -6
        assert transaction.reason == "optimization_start"
    finally:
        db.close()


def test_byok_mode_does_not_consume_credits(client, monkeypatch):
    from app.routes import optimization

    user_id, token = _create_user(credit_balance=0)
    monkeypatch.setattr(optimization, "BackgroundTasks", NoRunBackgroundTasks)
    monkeypatch.setattr(optimization, "run_optimization", _noop_run_optimization)
    _allow_public_model_url_dns(monkeypatch)

    response = client.post(
        "/api/optimization/start",
        json={
            "original_text": "test paragraph",
            "processing_mode": "paper_polish",
            "billing_mode": "byok",
            "polish_config": {
                "model": "gpt-5.4",
                "api_key": "sk-test",
                "base_url": "https://api.example/v1",
            },
        },
        headers={"Authorization": f"Bearer {token}"},
    )

    assert response.status_code == 200
    assert response.json()["charge_status"] == "not_charged"

    db = SessionLocal()
    try:
        user = db.query(User).filter(User.id == user_id).one()
        transactions = db.query(CreditTransaction).filter(CreditTransaction.user_id == user_id).all()
        assert user.credit_balance == 0
        assert transactions == []
    finally:
        db.close()


def test_byok_mode_rejects_private_request_base_url(client, monkeypatch):
    from app.routes import optimization

    _, token = _create_user(credit_balance=0)
    monkeypatch.setattr(optimization, "BackgroundTasks", NoRunBackgroundTasks)
    monkeypatch.setattr(optimization, "run_optimization", _noop_run_optimization)

    response = client.post(
        "/api/optimization/start",
        json={
            "original_text": "test paragraph",
            "processing_mode": "paper_polish",
            "billing_mode": "byok",
            "polish_config": {
                "model": "gpt-5.4",
                "api_key": "sk-test",
                "base_url": "https://127.0.0.1/v1",
            },
        },
        headers={"Authorization": f"Bearer {token}"},
    )

    assert response.status_code == 400
    assert "Base URL" in response.json()["detail"]


def test_platform_mode_rejects_user_with_insufficient_credits(client, monkeypatch):
    from app.routes import optimization

    _, token = _create_user(credit_balance=5)
    monkeypatch.setattr(optimization, "BackgroundTasks", NoRunBackgroundTasks)
    monkeypatch.setattr(optimization, "run_optimization", _noop_run_optimization)

    response = client.post(
        "/api/optimization/start",
        json={
            "original_text": "汉" * 3000,
            "processing_mode": "paper_polish_enhance",
            "billing_mode": "platform",
        },
        headers={"Authorization": f"Bearer {token}"},
    )

    assert response.status_code == 403
    assert response.json()["detail"] == "平台剩余啤酒不足，本次需要 6 啤酒，当前剩余 5 啤酒"


def test_ai_detect_reduce_start_does_not_hold_platform_credit(client, monkeypatch):
    from app.routes import optimization

    user_id, token = _create_user(credit_balance=0)
    monkeypatch.setattr(optimization, "BackgroundTasks", NoRunBackgroundTasks)
    monkeypatch.setattr(optimization, "run_optimization", _noop_run_optimization)
    monkeypatch.setattr(optimization, "zhuque_service", ReadyZhuqueService())

    response = client.post(
        "/api/optimization/start",
        json={
            "original_text": "汉" * 1000,
            "processing_mode": "ai_detect_reduce",
            "billing_mode": "platform",
        },
        headers={"Authorization": f"Bearer {token}"},
    )

    assert response.status_code == 200
    assert response.json()["processing_mode"] == "ai_detect_reduce"
    assert response.json()["current_stage"] == "ai_detect_reduce"
    assert response.json()["charge_status"] == "not_charged"
    assert response.json()["charged_credits"] == 0

    db = SessionLocal()
    try:
        user = db.query(User).filter(User.id == user_id).one()
        transactions = db.query(CreditTransaction).filter(CreditTransaction.user_id == user_id).all()
        assert user.credit_balance == 0
        assert transactions == []
    finally:
        db.close()


def test_ai_detect_reduce_retry_does_not_hold_platform_credit(client, monkeypatch):
    from app.routes import optimization

    user_id, token = _create_user(credit_balance=0)
    monkeypatch.setattr(optimization, "run_optimization", _noop_run_optimization)

    db = SessionLocal()
    try:
        failed_session = OptimizationSession(
            user_id=user_id,
            session_id="failed-ai-detect-reduce-no-prehold",
            original_text="汉" * 1000,
            current_stage="ai_detect_reduce",
            status="failed",
            error_message="朱雀降重未达标",
            processing_mode="ai_detect_reduce",
            billing_mode="platform",
            charge_status="not_charged",
            charged_credits=0,
        )
        db.add(failed_session)
        db.commit()
    finally:
        db.close()

    response = client.post(
        "/api/optimization/sessions/failed-ai-detect-reduce-no-prehold/retry",
        json={"billing_mode": "keep"},
        headers={"Authorization": f"Bearer {token}"},
    )

    assert response.status_code == 200

    db = SessionLocal()
    try:
        user = db.query(User).filter(User.id == user_id).one()
        session = db.query(OptimizationSession).filter(
            OptimizationSession.session_id == "failed-ai-detect-reduce-no-prehold"
        ).one()
        transactions = db.query(CreditTransaction).filter(CreditTransaction.user_id == user_id).all()
        assert user.credit_balance == 0
        assert session.status == "queued"
        assert session.billing_mode == "platform"
        assert session.charge_status == "not_charged"
        assert session.charged_credits == 0
        assert transactions == []
    finally:
        db.close()


def test_ai_detect_reduce_retry_clears_stale_zhuque_trace_final(client, monkeypatch):
    from app.routes import optimization

    user_id, token = _create_user(credit_balance=0)
    monkeypatch.setattr(optimization, "run_optimization", _noop_run_optimization)

    db = SessionLocal()
    try:
        failed_session = OptimizationSession(
            user_id=user_id,
            session_id="failed-ai-detect-reduce-stale-final",
            original_text="汉" * 1000,
            current_stage="ai_detect_reduce",
            status="failed",
            error_message="朱雀降重未达标",
            processing_mode="ai_detect_reduce",
            billing_mode="platform",
            charge_status="not_charged",
            charged_credits=0,
            zhuque_agent_trace=json.dumps(
                {
                    "version": 1,
                    "threshold": 20.0,
                    "events": [
                        {
                            "type": "detect",
                            "seq": 1,
                            "message": "上一轮检测事件保留",
                        }
                    ],
                    "final": {
                        "status": "failed",
                        "rate": None,
                        "diagnosis": "上一轮失败诊断不应在重试中继续展示",
                    },
                },
                ensure_ascii=False,
            ),
        )
        db.add(failed_session)
        db.commit()
    finally:
        db.close()

    response = client.post(
        "/api/optimization/sessions/failed-ai-detect-reduce-stale-final/retry",
        json={"billing_mode": "keep"},
        headers={"Authorization": f"Bearer {token}"},
    )

    assert response.status_code == 200

    db = SessionLocal()
    try:
        session = db.query(OptimizationSession).filter(
            OptimizationSession.session_id == "failed-ai-detect-reduce-stale-final"
        ).one()
        trace = json.loads(session.zhuque_agent_trace)
        assert session.status == "queued"
        assert "final" not in trace
        assert trace["events"][0]["message"] == "上一轮检测事件保留"
        assert trace["threshold"] == 20.0
    finally:
        db.close()


def test_session_stream_uses_short_lived_stream_token_instead_of_login_token(client):
    user_id, token = _create_user(credit_balance=0)
    db = SessionLocal()
    try:
        session = OptimizationSession(
            user_id=user_id,
            session_id="stream-token-session",
            original_text="测试正文",
            current_stage="polish",
            status="queued",
            progress=0,
            processing_mode="paper_polish",
            billing_mode="platform",
            charge_status="not_charged",
        )
        db.add(session)
        db.commit()
    finally:
        db.close()

    progress_with_query_login_token = client.get(
        "/api/optimization/sessions/stream-token-session/progress",
        params={"access_token": token},
    )
    assert progress_with_query_login_token.status_code == 401

    stream_with_query_login_token = client.get(
        "/api/optimization/sessions/stream-token-session/stream",
        params={"access_token": token},
    )
    assert stream_with_query_login_token.status_code == 401

    token_response = client.post(
        "/api/optimization/sessions/stream-token-session/stream-token",
        headers={"Authorization": f"Bearer {token}"},
    )
    assert token_response.status_code == 200
    body = token_response.json()
    assert body["token_type"] == "bearer"
    assert body["expires_in"] == config_module.settings.STREAM_TOKEN_EXPIRE_SECONDS
    assert body["stream_token"]
    assert body["stream_token"] != token

    payload = verify_stream_token(body["stream_token"])
    assert payload["sub"] == str(user_id)
    assert payload["role"] == "stream"
    assert payload["scope"] == "session_stream"
    assert payload["session_id"] == "stream-token-session"
    assert payload["token_version"] == 0

    stream_with_login_token_in_stream_param = client.get(
        "/api/optimization/sessions/stream-token-session/stream",
        params={"stream_token": token},
    )
    assert stream_with_login_token_in_stream_param.status_code == 401


def test_credit_service_public_error_messages_use_beer_unit():
    from fastapi import HTTPException
    from app.services.credit_service import CreditService

    user_id, _ = _create_user(credit_balance=0)
    db = SessionLocal()
    try:
        user = db.query(User).filter(User.id == user_id).one()
        service = CreditService(db)

        error_checks = [
            (lambda: service.hold_platform_credit(user, reason="test", amount=0), "扣除啤酒必须大于 0"),
            (lambda: service.refund_platform_credit(user, reason="test", amount=0), "退回啤酒必须大于 0"),
            (lambda: service.add_credits(user, amount=0, reason="test"), "充值啤酒必须大于 0"),
        ]

        for action, expected_detail in error_checks:
            try:
                action()
            except HTTPException as exc:
                assert exc.detail == expected_detail
            else:
                raise AssertionError("expected HTTPException")
    finally:
        db.close()


def test_failed_platform_job_refunds_held_credits_once():
    from app.services.credit_service import CreditService

    user_id, _ = _create_user(credit_balance=5)
    db = SessionLocal()
    try:
        user = db.query(User).filter(User.id == user_id).one()
        session = OptimizationSession(
            user_id=user.id,
            session_id="session-1",
            original_text="test paragraph",
            current_stage="polish",
            status="queued",
            billing_mode="platform",
            charge_status="held",
            charged_credits=3,
        )
        db.add(session)
        db.flush()
        CreditService(db).hold_platform_credit(user, reason="optimization_start", session_id=session.id, amount=3)
        db.commit()

        CreditService(db).refund_held_platform_credit(session)
        db.commit()
        CreditService(db).refund_held_platform_credit(session)
        db.commit()

        db.refresh(user)
        assert user.credit_balance == 5
        assert session.charge_status == "refunded"
        transactions = db.query(CreditTransaction).filter(CreditTransaction.user_id == user.id).all()
        assert [txn.delta for txn in transactions] == [-3, 3]
    finally:
        db.close()
