from fastapi import APIRouter, Depends, HTTPException, BackgroundTasks, Request
from sqlalchemy.orm import Session, defer, joinedload
from sqlalchemy import func, and_, case
from typing import List, Optional
import base64
import importlib.util
import os
import subprocess
import sys
from io import BytesIO
from pathlib import Path
import json
import re
from app.database import get_db
from app.models.models import User, OptimizationSession, OptimizationSegment, ChangeLog, PaperProject
from app.schemas import (
    OptimizationCreate, SessionResponse, SessionDetailResponse,
    QueueStatusResponse, ProgressUpdate, ChangeLogResponse, ExportConfirmation,
    SessionRetryRequest, SessionProjectUpdateRequest, ZhuqueBrowserLaunchResponse, ZhuqueBrowserStatusResponse,
    ZhuquePreflightRequest, ZhuqueReadinessResponse,
)
from app.services.concurrency import concurrency_manager
from app.services.credit_service import CreditService, calculate_optimization_credits
from app.services.provider_config_service import ProviderConfigService
from app.services.stream_manager import stream_manager
from app.services.task_queue import process_session_by_id
from app.services.zhuque_service import zhuque_service, zhuque_user_dir
from app.services.zhuque_remote_login_service import zhuque_remote_login_service
from app.services.ai_service import count_text_length, split_text_into_segments
from app.utils.auth import generate_session_id, get_current_user_with_legacy_fallback
from app.utils.url_security import validate_model_base_url
from app.utils.time import utcnow
from datetime import datetime, timedelta
import asyncio
from app.config import settings
from sse_starlette.sse import EventSourceResponse
from docx import Document

router = APIRouter(prefix="/optimization", tags=["optimization"])

ONLINE_USER_WINDOW_SECONDS = 60


def _clean_export_filename_part(value: str, fallback: str) -> str:
    cleaned = re.sub(r'[\\/:*?"<>|\r\n\t]+', "_", (value or "").strip())
    cleaned = re.sub(r"\s+", " ", cleaned).strip(" ._")
    return cleaned or fallback


def _build_export_filename(session: OptimizationSession, extension: str) -> str:
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    project_title = session.project.title if session.project and session.project.title else ""
    project_part = _clean_export_filename_part(project_title, "未归档")

    parts = [project_part]
    if project_title and session.task_title and session.task_title.strip():
        parts.append(_clean_export_filename_part(session.task_title, "本次处理"))
    parts.append(timestamp)
    return f"{'_'.join(parts)}.{extension}"


def _build_aigc_report_filename(session: OptimizationSession, extension: str) -> str:
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    project_title = session.project.title if session.project and session.project.title else ""
    project_part = _clean_export_filename_part(project_title, "未归档")

    parts = [project_part]
    if project_title and session.task_title and session.task_title.strip():
        parts.append(_clean_export_filename_part(session.task_title, "本次处理"))
    parts.extend(["AIGC检测报告", timestamp])
    return f"{'_'.join(parts)}.{extension}"


def _build_docx_base64(text: str) -> str:
    document = Document()
    paragraphs = text.split("\n\n") if text else [""]
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)

    buffer = BytesIO()
    document.save(buffer)
    return base64.b64encode(buffer.getvalue()).decode("ascii")


def _parse_zhuque_result(raw_result: str | dict | None) -> dict:
    if isinstance(raw_result, dict):
        return raw_result
    if not raw_result:
        return {}
    try:
        parsed = json.loads(raw_result)
    except (TypeError, json.JSONDecodeError):
        return {"raw": str(raw_result)}
    return parsed if isinstance(parsed, dict) else {"raw": parsed}


def _zhuque_rate_percent(value) -> Optional[float]:
    try:
        number = float(value)
    except (TypeError, ValueError):
        return None
    return round(max(0.0, min(number, 100.0)), 2)


def _zhuque_ratio_percent(value) -> Optional[float]:
    try:
        number = float(value)
    except (TypeError, ValueError):
        return None
    # 朱雀 labels_ratio 是 0-1；转换成报告里的 0-100。
    if 0 <= number <= 1:
        number *= 100
    return _zhuque_rate_percent(number)


def _zhuque_risk_rate_from_result(result: dict, fallback: Optional[float] = None) -> Optional[float]:
    if not isinstance(result, dict) or result.get("success") is False:
        return None

    labels_ratio = result.get("labels_ratio") or {}
    if isinstance(labels_ratio, dict) and labels_ratio:
        ai_rate = _zhuque_ratio_from_result(result, "0") or 0.0
        suspicious_rate = _zhuque_ratio_from_result(result, "2") or 0.0
        return round(max(ai_rate, suspicious_rate), 2)

    for key in ("risk_rate", "rate"):
        rate = _zhuque_rate_percent(result.get(key))
        if rate is not None:
            return rate
    return _zhuque_rate_percent(fallback)


def _zhuque_ratio_from_result(result: dict, label: str) -> Optional[float]:
    labels_ratio = result.get("labels_ratio") or {}
    if not isinstance(labels_ratio, dict):
        return None
    legacy_keys = {
        "0": ("ai", "AI", "ai_rate", "aiGenerated"),
        "1": ("human", "Human", "human_rate"),
        "2": ("suspicious", "mixed", "疑似AI"),
    }
    candidates = [label]
    if label.isdigit():
        candidates.append(int(label))
    candidates.extend(legacy_keys.get(label, ()))
    for key in candidates:
        if key in labels_ratio:
            return _zhuque_ratio_percent(labels_ratio.get(key))
    return None


def _format_report_rate(rate: Optional[float]) -> str:
    if rate is None:
        return "--"
    return f"{rate:.1f}%"


def _format_report_number(value) -> str:
    try:
        number = int(value)
    except (TypeError, ValueError):
        return "--"
    return "--" if number < 0 else str(number)


def _safe_report_cell(value: object) -> str:
    text = "" if value is None else str(value)
    return text.replace("\r", " ").replace("\n", " ").strip()


def _build_joined_report_spans(segments: List[OptimizationSegment]) -> List[tuple[OptimizationSegment, str, int, int]]:
    spans: List[tuple[OptimizationSegment, str, int, int]] = []
    cursor = 0
    ordered_segments = sorted(segments, key=lambda item: item.segment_index)
    for index, seg in enumerate(ordered_segments):
        text = seg.zhuque_reduced_text or seg.enhanced_text or seg.polished_text or seg.original_text or ""
        start = cursor
        end = start + len(text)
        spans.append((seg, text, start, end))
        cursor = end
        if index < len(ordered_segments) - 1:
            cursor += 2
    return spans


def _segment_report_rows_from_labels(
    segments: List[OptimizationSegment],
    zhuque_result: dict,
    fallback_rate: Optional[float],
) -> List[dict]:
    spans = _build_joined_report_spans(segments)
    labels = zhuque_result.get("segment_labels") or []
    usable_labels = []
    if isinstance(labels, list):
        for item in labels:
            if not isinstance(item, dict):
                continue
            position = item.get("position")
            if (
                not isinstance(position, list)
                or len(position) != 2
                or not all(isinstance(value, (int, float)) for value in position)
            ):
                continue
            start, end = int(position[0]), int(position[1])
            if end <= start:
                continue
            try:
                label = int(item.get("label"))
            except (TypeError, ValueError):
                continue
            usable_labels.append({"label": label, "start": start, "end": end})

    rows: List[dict] = []
    full_ai_rate = _zhuque_ratio_from_result(zhuque_result, "0")
    full_human_rate = _zhuque_ratio_from_result(zhuque_result, "1")
    full_suspicious_rate = _zhuque_ratio_from_result(zhuque_result, "2")
    full_risk_rate = _zhuque_risk_rate_from_result(zhuque_result, fallback_rate)

    for seg, text, seg_start, seg_end in spans:
        text_span = max(seg_end - seg_start, 0)
        label_chars = {0: 0, 1: 0, 2: 0}
        for item in usable_labels:
            overlap = max(0, min(seg_end, item["end"]) - max(seg_start, item["start"]))
            if overlap > 0 and item["label"] in label_chars:
                label_chars[item["label"]] += overlap

        if usable_labels and text_span > 0:
            denominator = max(text_span, 1)
            ai_rate = round(label_chars[0] / denominator * 100, 2)
            human_rate = round(label_chars[1] / denominator * 100, 2)
            suspicious_rate = round(label_chars[2] / denominator * 100, 2)
            risk_rate = round(min(ai_rate + suspicious_rate, 100.0), 2)
            source = "segment_labels"
        else:
            ai_rate = full_ai_rate
            human_rate = full_human_rate
            suspicious_rate = full_suspicious_rate
            risk_rate = _zhuque_rate_percent(seg.zhuque_detect_rate) or full_risk_rate
            source = "full_text_fallback"

        rows.append({
            "segment_index": seg.segment_index,
            "text": text,
            "char_count": count_text_length(text),
            "ai_rate": ai_rate,
            "human_rate": human_rate,
            "suspicious_rate": suspicious_rate,
            "risk_rate": risk_rate,
            "source": source,
            "status": "高风险" if (risk_rate is not None and risk_rate >= 50) else "需关注" if (risk_rate is not None and risk_rate >= 20) else "低风险",
        })
    return rows


def _build_aigc_report_payload(session: OptimizationSession, segments: List[OptimizationSegment]) -> dict:
    ordered_segments = sorted(segments, key=lambda item: item.segment_index)
    result_segment = next(
        (
            seg for seg in reversed(ordered_segments)
            if seg.zhuque_detect_result or seg.zhuque_detect_rate is not None
        ),
        ordered_segments[-1] if ordered_segments else None,
    )
    zhuque_result = _parse_zhuque_result(result_segment.zhuque_detect_result if result_segment else None)
    fallback_rate = result_segment.zhuque_detect_rate if result_segment else None
    final_risk_rate = _zhuque_risk_rate_from_result(zhuque_result, fallback_rate)
    rows = _segment_report_rows_from_labels(ordered_segments, zhuque_result, fallback_rate)
    threshold = float(settings.ZHUQUE_DETECT_THRESHOLD)
    high_risk_count = sum(1 for row in rows if row["risk_rate"] is not None and row["risk_rate"] > threshold)

    return {
        "title": "GankAIGC AIGC 检测报告",
        "session_id": session.session_id,
        "project_title": session.project.title if session.project and session.project.title else "未归档",
        "task_title": session.task_title or "",
        "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "threshold": threshold,
        "final_risk_rate": final_risk_rate,
        "ai_rate": _zhuque_ratio_from_result(zhuque_result, "0"),
        "human_rate": _zhuque_ratio_from_result(zhuque_result, "1"),
        "suspicious_rate": _zhuque_ratio_from_result(zhuque_result, "2"),
        "detect_count": max((seg.zhuque_detect_count or 0) for seg in ordered_segments) if ordered_segments else 0,
        "reduce_rounds": max((seg.zhuque_reduce_attempt or 0) for seg in ordered_segments) if ordered_segments else 0,
        "remaining_uses": zhuque_result.get("remaining_uses"),
        "text_length": zhuque_result.get("text_length"),
        "message": zhuque_result.get("message") or zhuque_result.get("alert_text") or "",
        "segment_count": len(rows),
        "high_risk_count": high_risk_count,
        "rows": rows,
        "source": zhuque_result.get("source") or "zhuque",
    }


def _build_aigc_report_markdown(payload: dict) -> str:
    lines = [
        f"# {payload['title']}",
        "",
        "## 报告摘要",
        "",
        f"- 会话 ID：{payload['session_id']}",
        f"- 项目：{payload['project_title']}",
        f"- 任务：{payload['task_title'] or '--'}",
        f"- 生成时间：{payload['generated_at']}",
        f"- 最终风险率：{_format_report_rate(payload['final_risk_rate'])}",
        f"- AI特征：{_format_report_rate(payload['ai_rate'])}",
        f"- 疑似AI：{_format_report_rate(payload['suspicious_rate'])}",
        f"- 人工特征：{_format_report_rate(payload['human_rate'])}",
        f"- 阈值：{_format_report_rate(payload['threshold'])}",
        f"- 朱雀检测次数：{payload['detect_count']} 次",
        f"- 降重轮次：{payload['reduce_rounds']} 轮",
        f"- 朱雀剩余次数：{_format_report_number(payload['remaining_uses'])}",
        f"- 高风险段落：{payload['high_risk_count']} / {payload['segment_count']}",
    ]
    if payload.get("message"):
        lines.append(f"- 朱雀提示：{payload['message']}")

    lines.extend([
        "",
        "## 逐段 AI 率",
        "",
        "| 段落 | 字数 | 段落AI率 | AI特征 | 疑似AI | 人工特征 | 结论 |",
        "| --- | ---: | ---: | ---: | ---: | ---: | --- |",
    ])
    for row in payload["rows"]:
        lines.append(
            "| "
            f"{row['segment_index'] + 1} | "
            f"{row['char_count']} | "
            f"{_format_report_rate(row['risk_rate'])} | "
            f"{_format_report_rate(row['ai_rate'])} | "
            f"{_format_report_rate(row['suspicious_rate'])} | "
            f"{_format_report_rate(row['human_rate'])} | "
            f"{row['status']} |"
        )

    lines.extend(["", "## 段落明细", ""])
    for row in payload["rows"]:
        lines.extend([
            f"### 段落 {row['segment_index'] + 1}｜AI率 {_format_report_rate(row['risk_rate'])}",
            "",
            f"- AI特征：{_format_report_rate(row['ai_rate'])}",
            f"- 疑似AI：{_format_report_rate(row['suspicious_rate'])}",
            f"- 人工特征：{_format_report_rate(row['human_rate'])}",
            f"- 结论：{row['status']}",
            "",
            row["text"] or "（空段落）",
            "",
        ])
    return "\n".join(lines)


def _build_aigc_report_docx_base64(payload: dict) -> str:
    document = Document()
    document.add_heading(payload["title"], level=0)
    document.add_paragraph(
        "本报告基于朱雀 AI 检测结果生成，按最终导出文本映射每一段的 AI 特征、疑似 AI 与人工特征占比。"
    )

    summary_table = document.add_table(rows=0, cols=2)
    summary_items = [
        ("会话 ID", payload["session_id"]),
        ("项目", payload["project_title"]),
        ("任务", payload["task_title"] or "--"),
        ("生成时间", payload["generated_at"]),
        ("最终风险率", _format_report_rate(payload["final_risk_rate"])),
        ("AI特征", _format_report_rate(payload["ai_rate"])),
        ("疑似AI", _format_report_rate(payload["suspicious_rate"])),
        ("人工特征", _format_report_rate(payload["human_rate"])),
        ("阈值", _format_report_rate(payload["threshold"])),
        ("朱雀检测次数", f"{payload['detect_count']} 次"),
        ("降重轮次", f"{payload['reduce_rounds']} 轮"),
        ("朱雀剩余次数", _format_report_number(payload["remaining_uses"])),
        ("高风险段落", f"{payload['high_risk_count']} / {payload['segment_count']}"),
    ]
    if payload.get("message"):
        summary_items.append(("朱雀提示", payload["message"]))
    for key, value in summary_items:
        row = summary_table.add_row().cells
        row[0].text = str(key)
        row[1].text = _safe_report_cell(value)

    document.add_heading("逐段 AI 率", level=1)
    table = document.add_table(rows=1, cols=7)
    headers = ["段落", "字数", "段落AI率", "AI特征", "疑似AI", "人工特征", "结论"]
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for row_data in payload["rows"]:
        cells = table.add_row().cells
        cells[0].text = str(row_data["segment_index"] + 1)
        cells[1].text = str(row_data["char_count"])
        cells[2].text = _format_report_rate(row_data["risk_rate"])
        cells[3].text = _format_report_rate(row_data["ai_rate"])
        cells[4].text = _format_report_rate(row_data["suspicious_rate"])
        cells[5].text = _format_report_rate(row_data["human_rate"])
        cells[6].text = row_data["status"]

    document.add_heading("段落明细", level=1)
    for row_data in payload["rows"]:
        document.add_heading(
            f"段落 {row_data['segment_index'] + 1}｜AI率 {_format_report_rate(row_data['risk_rate'])}",
            level=2,
        )
        document.add_paragraph(
            "AI特征 {ai}；疑似AI {suspicious}；人工特征 {human}；结论：{status}".format(
                ai=_format_report_rate(row_data["ai_rate"]),
                suspicious=_format_report_rate(row_data["suspicious_rate"]),
                human=_format_report_rate(row_data["human_rate"]),
                status=row_data["status"],
            )
        )
        document.add_paragraph(row_data["text"] or "（空段落）")

    buffer = BytesIO()
    document.save(buffer)
    return base64.b64encode(buffer.getvalue()).decode("ascii")


async def run_optimization(session_id: int):
    """后台运行已入队的优化任务。"""
    await process_session_by_id(session_id)


def _clear_session_provider_fields(session: OptimizationSession) -> None:
    session.polish_model = None
    session.polish_api_key = None
    session.polish_base_url = None
    session.enhance_model = None
    session.enhance_api_key = None
    session.enhance_base_url = None
    session.emotion_model = None
    session.emotion_api_key = None
    session.emotion_base_url = None


def _apply_retry_billing_mode(
    *,
    session: OptimizationSession,
    user: User,
    requested_billing_mode: str,
    db: Session,
) -> None:
    """重试失败任务时按用户当前选择的计费/API 模式刷新会话运行配置。"""
    target_billing_mode = session.billing_mode if requested_billing_mode == "keep" else requested_billing_mode

    if target_billing_mode == "byok":
        provider_config = ProviderConfigService(db).get_runtime_config(user)
        CreditService(db).refund_held_platform_credit(session)

        session.billing_mode = "byok"
        session.credential_source = "user_saved"
        session.charge_status = "not_charged"
        session.charged_credits = 0
        session.polish_model = provider_config["polish_model"]
        session.polish_api_key = None
        session.polish_base_url = provider_config["base_url"]
        session.enhance_model = provider_config["enhance_model"]
        session.enhance_api_key = None
        session.enhance_base_url = provider_config["base_url"]
        session.emotion_model = provider_config["emotion_model"]
        session.emotion_api_key = None
        session.emotion_base_url = provider_config["base_url"] if provider_config["emotion_model"] else None
        return

    if target_billing_mode == "platform":
        if session.processing_mode == "ai_detect_reduce":
            session.billing_mode = "platform"
            session.credential_source = "system"
            session.charge_status = "not_charged"
            session.charged_credits = 0
            _clear_session_provider_fields(session)
            return

        already_held = (
            session.billing_mode == "platform"
            and session.charge_status == "held"
        )
        required_credits = calculate_optimization_credits(session.original_text, session.processing_mode)

        session.billing_mode = "platform"
        session.credential_source = "system"
        _clear_session_provider_fields(session)

        if not already_held:
            CreditService(db).hold_platform_credit(
                user,
                reason="optimization_start",
                session_id=session.id,
                amount=required_credits,
            )
            session.charge_status = "held"
            session.charged_credits = 0 if user.is_unlimited else required_credits


def _validate_request_model_base_url(config, label: str) -> str:
    if not config or not config.base_url:
        raise HTTPException(status_code=400, detail=f"{label} Base URL 未配置")
    try:
        return validate_model_base_url(config.base_url)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc






def _zhuque_service_for_user(user: User):
    """Return the Zhuque service scoped to the current user; keeps old tests compatible."""
    for_user = getattr(zhuque_service, "for_user", None)
    return for_user(user.id) if callable(for_user) else zhuque_service


def _zhuque_capture_script_path() -> Optional[Path]:
    here = Path(__file__).resolve()
    candidates = [
        here.parents[4] / "zhuque_pkg" / "capture_zhuque_creds.py",
        here.parents[3] / "zhuque_pkg" / "capture_zhuque_creds.py",
        Path.cwd() / "zhuque_pkg" / "capture_zhuque_creds.py",
        Path.cwd().parent / "zhuque_pkg" / "capture_zhuque_creds.py",
    ]
    for candidate in candidates:
        if candidate.exists():
            return candidate
    return None


def _zhuque_playwright_browsers_path() -> Path:
    return Path(__file__).resolve().parents[3] / ".playwright-browsers"


def _zhuque_is_wsl() -> bool:
    if os.environ.get("WSL_INTEROP") or os.environ.get("WSL_DISTRO_NAME"):
        return True
    try:
        release = Path("/proc/sys/kernel/osrelease").read_text(encoding="utf-8").lower()
    except OSError:
        return False
    return "microsoft" in release or "wsl" in release


def _zhuque_windows_to_wsl_path(path: str) -> Optional[Path]:
    text = (path or "").strip().strip('"')
    if not text:
        return None
    try:
        completed = subprocess.run(
            ["wslpath", "-u", text],
            capture_output=True,
            text=True,
            timeout=2,
            check=False,
        )
    except (OSError, subprocess.SubprocessError):
        completed = None
    if completed and completed.stdout.strip():
        return Path(completed.stdout.strip())
    match = re.match(r"^([a-zA-Z]):\\(.*)$", text)
    if match:
        drive, rest = match.groups()
        linux_rest = rest.replace("\\", "/")
        return Path(f"/mnt/{drive.lower()}/{linux_rest}")
    return None


def _zhuque_local_browser_executable() -> Optional[Path]:
    """Return a Chromium-family browser executable Playwright can control.

    On WSL, prefer Windows Chrome/Edge/Brave. The capture script launches it
    with a dedicated CDP port and connects from Playwright, so the visible scan
    window appears on the user's Windows desktop instead of inside WSL.
    """
    env_path = os.environ.get("ZHUQUE_CHROME_EXECUTABLE")
    windows_env_path = _zhuque_windows_to_wsl_path(env_path) if _zhuque_is_wsl() and env_path else None
    candidates = [
        env_path,
        windows_env_path,
        "/usr/bin/google-chrome",
        "/usr/bin/google-chrome-stable",
        "/usr/bin/chromium",
        "/usr/bin/chromium-browser",
        "/usr/bin/microsoft-edge",
        "/usr/bin/microsoft-edge-stable",
        "/usr/bin/msedge",
        "/usr/bin/brave-browser",
        "/usr/bin/brave",
        "/snap/bin/chromium",
        "/snap/bin/brave",
    ]
    if _zhuque_is_wsl():
        candidates.extend(
            [
                "/mnt/c/Program Files/Google/Chrome/Application/chrome.exe",
                "/mnt/c/Program Files (x86)/Google/Chrome/Application/chrome.exe",
                "/mnt/c/Program Files/Microsoft/Edge/Application/msedge.exe",
                "/mnt/c/Program Files (x86)/Microsoft/Edge/Application/msedge.exe",
                "/mnt/c/Program Files/BraveSoftware/Brave-Browser/Application/brave.exe",
                "/mnt/c/Program Files (x86)/BraveSoftware/Brave-Browser/Application/brave.exe",
            ]
        )
    for candidate in candidates:
        if candidate and Path(candidate).exists():
            return Path(candidate)
    return None


def _zhuque_playwright_browser_ready() -> bool:
    if _zhuque_local_browser_executable() is not None:
        return True
    browser_root = Path(os.environ.get("PLAYWRIGHT_BROWSERS_PATH") or _zhuque_playwright_browsers_path())
    executable_patterns = [
        "*/chrome-linux/chrome",
        "*/chrome-linux64/chrome",
        "*/chromium*/chrome",
    ]
    for pattern in executable_patterns:
        if any(candidate.exists() for candidate in browser_root.glob(pattern)):
            return True
    return False


def _zhuque_capture_env(user: Optional[User] = None) -> dict:
    env = os.environ.copy()
    env.setdefault("PLAYWRIGHT_BROWSERS_PATH", str(_zhuque_playwright_browsers_path()))
    env.setdefault("ZHUQUE_CDP_PORT", str(settings.ZHUQUE_CDP_PORT))
    if user is not None:
        capture_dir = zhuque_user_dir(user.id)
        capture_dir.mkdir(parents=True, exist_ok=True)
        env["ZHUQUE_CAPTURE_DIR"] = str(capture_dir)
    browser_executable = _zhuque_local_browser_executable()
    if browser_executable is not None:
        env.setdefault("ZHUQUE_CHROME_EXECUTABLE", str(browser_executable))
    return env


def _start_zhuque_wechat_capture(*, sync_session: bool = True, user: Optional[User] = None) -> dict:
    service = _zhuque_service_for_user(user) if user is not None else zhuque_service
    status = service._ensure_api().credential_status()
    script_path = _zhuque_capture_script_path()
    if script_path is None:
        return {
            "status": "missing_script",
            "auth_mode": "headless_api",
            "login_mode": "wechat_qr",
            "credential_file": status.get("credential_file", ""),
            "sync_session": sync_session,
            "command": "python zhuque_pkg/capture_zhuque_creds.py --sync-session",
            "message": "未找到 zhuque_pkg/capture_zhuque_creds.py，请确认新朱雀包在项目根目录",
        }

    launch_args = ["--sync-session"] if sync_session else []
    command = " ".join([sys.executable, str(script_path), *launch_args]).strip()
    if importlib.util.find_spec("playwright") is None:
        browsers_path = _zhuque_playwright_browsers_path()
        return {
            "status": "manual_required",
            "auth_mode": "headless_api",
            "login_mode": "wechat_qr",
            "credential_file": status.get("credential_file", str(script_path.parent / "creds_latest.json")),
            "sync_session": sync_session,
            "command": f'{sys.executable} -m pip install playwright && PLAYWRIGHT_BROWSERS_PATH="{browsers_path}" {sys.executable} -m playwright install chromium && {command}',
            "message": "当前 Python 环境未安装 Playwright，无法自动打开朱雀扫码授权页/真实网页状态同步窗口。请先安装 Playwright 并执行同步命令。",
        }
    if not _zhuque_playwright_browser_ready():
        browsers_path = _zhuque_playwright_browsers_path()
        return {
            "status": "manual_required",
            "auth_mode": "headless_api",
            "login_mode": "wechat_qr",
            "credential_file": status.get("credential_file", str(script_path.parent / "creds_latest.json")),
            "sync_session": sync_session,
            "command": f'PLAYWRIGHT_BROWSERS_PATH="{browsers_path}" {sys.executable} -m playwright install chromium && {command}',
            "message": "Playwright 已安装，但未找到可用于朱雀状态同步窗口的 Chromium 内核浏览器（Chrome/Chromium/Edge/Brave）或 Playwright 内置 Chromium。请先安装任一浏览器内核；登录/退出状态只按朱雀网页真实状态同步。",
        }

    try:
        log_dir = zhuque_user_dir(user.id) if user is not None else script_path.parent
        log_dir.mkdir(parents=True, exist_ok=True)
        log_path = log_dir / "capture_latest.log"
        log_handle = open(log_path, "a", encoding="utf-8", buffering=1)
        log_handle.write(f"\n--- start {utcnow().isoformat()} sync_session={sync_session} ---\n")
        subprocess.Popen(
            [sys.executable, str(script_path), *launch_args],
            cwd=str(script_path.parent),
            env=_zhuque_capture_env(user),
            stdout=log_handle,
            stderr=subprocess.STDOUT,
            start_new_session=True,
        )
        return {
            "status": "started",
            "auth_mode": "headless_api",
            "login_mode": "wechat_qr",
            "credential_file": status.get("credential_file", str(script_path.parent / "creds_latest.json")),
            "sync_session": sync_session,
            "command": command,
            "message": (
                "已打开朱雀真实网页状态同步窗口；请在朱雀网页内登录/退出，GankAIGC 会按网页状态保存或清除凭证。"
                if sync_session
                else "已打开朱雀微信扫码授权页；扫码完成后会保存凭证，后续检测走无头 API。"
            ) + f"关闭窗口后会保留最后一次同步状态。日志: {log_path}",
        }
    except Exception as exc:
        return {
            "status": "manual_required",
            "auth_mode": "headless_api",
            "login_mode": "wechat_qr",
            "credential_file": status.get("credential_file", str(script_path.parent / "creds_latest.json")),
            "sync_session": sync_session,
            "command": command,
            "message": f"自动打开朱雀网页状态同步窗口失败，请在终端手动运行命令完成登录态同步：{command}；错误: {exc}",
        }


def _get_zhuque_headless_status(user: Optional[User] = None) -> dict:
    service = _zhuque_service_for_user(user) if user is not None else zhuque_service
    status = service._ensure_api().credential_status()
    ready = bool(status.get("ready"))
    return {
        "status": "connected" if ready else "missing_credentials",
        "connected": ready,
        "ready": ready,
        "has_token": bool(status.get("has_token")),
        "remaining_uses": status.get("remaining_uses", -1),
        "button_enabled": bool(status.get("button_enabled", ready)),
        "auth_mode": "headless_api",
        "login_mode": "wechat_qr",
        "credential_file": status.get("credential_file", ""),
        "user_name": status.get("user_name", ""),
        "quota_text": status.get("quota_text", ""),
        "captured_at": status.get("captured_at", ""),
        "message": status.get("message") or ("朱雀无头 API 已就绪" if ready else "未找到朱雀微信扫码凭证"),
    }

def _with_zhuque_cost_estimate(readiness: dict, text: str) -> dict:
    segments = split_text_into_segments(text or "") if text else []
    segment_count = max(len(segments), 1) if text else 0
    return {
        **readiness,
        "estimated_first_round_credits": segment_count * 10,
        "estimated_max_round_credits": segment_count * settings.ZHUQUE_MAX_REDUCE_ROUNDS * 10,
    }


def _zhuque_preflight_error(readiness: dict) -> str:
    message = readiness.get("message") or "朱雀尚未就绪"
    actions = readiness.get("actions") or []
    if actions:
        return f"{message}；请先" + "、".join(actions)
    return message


async def _run_zhuque_preflight_or_raise(text: str, user: User) -> dict:
    if len(text or "") < 350:
        raise HTTPException(status_code=400, detail=f"朱雀 AI 检测要求文本长度不少于 350 字，当前 {len(text or '')} 字")
    readiness = _with_zhuque_cost_estimate(await _zhuque_service_for_user(user).readiness(text), text)
    if not readiness.get("ready"):
        raise HTTPException(status_code=400, detail=_zhuque_preflight_error(readiness))
    return readiness


@router.post("/start", response_model=SessionResponse)
async def start_optimization(
    data: OptimizationCreate,
    background_tasks: BackgroundTasks,
    user: User = Depends(get_current_user_with_legacy_fallback),
    db: Session = Depends(get_db)
):
    """开始优化任务"""
    usage_count = user.usage_count or 0
    
    # 验证处理模式
    valid_modes = ['paper_polish', 'paper_enhance', 'paper_polish_enhance', 'emotion_polish', 'ai_detect_reduce']
    if data.processing_mode not in valid_modes:
        raise HTTPException(
            status_code=400,
            detail=f"无效的处理模式。支持的模式: {', '.join(valid_modes)}"
        )
    required_credits = calculate_optimization_credits(data.original_text, data.processing_mode)

    project = None
    if data.project_id is not None:
        project = (
            db.query(PaperProject)
            .filter(
                PaperProject.id == data.project_id,
                PaperProject.user_id == user.id,
                PaperProject.is_archived.is_(False),
            )
            .first()
        )
        if not project:
            raise HTTPException(status_code=404, detail="论文项目不存在")

    # 根据处理模式设置初始阶段
    if data.processing_mode == 'emotion_polish':
        initial_stage = 'emotion_polish'
    elif data.processing_mode == 'ai_detect_reduce':
        initial_stage = 'ai_detect_reduce'
    elif data.processing_mode == 'paper_enhance':
        initial_stage = 'enhance'
    else:
        initial_stage = 'polish'
    
    provider_config = None
    request_polish_base_url = None
    request_enhance_base_url = None
    request_emotion_base_url = None
    if data.billing_mode == "byok":
        if data.polish_config:
            request_polish_base_url = _validate_request_model_base_url(data.polish_config, "润色模型")
            request_enhance_base_url = (
                _validate_request_model_base_url(data.enhance_config, "增强模型")
                if data.enhance_config
                else request_polish_base_url
            )
            request_emotion_base_url = (
                _validate_request_model_base_url(data.emotion_config, "感情润色模型")
                if data.emotion_config
                else None
            )
            provider_config = {
                "base_url": request_polish_base_url,
                "api_key": data.polish_config.api_key,
                "polish_model": data.polish_config.model,
                "enhance_model": data.enhance_config.model if data.enhance_config else data.polish_config.model,
                "emotion_model": data.emotion_config.model if data.emotion_config else None,
            }
        else:
            provider_config = ProviderConfigService(db).get_runtime_config(user)

    if data.processing_mode == "ai_detect_reduce":
        await _run_zhuque_preflight_or_raise(data.original_text, user)

    polish_model = data.polish_config.model if data.polish_config else None
    polish_api_key = data.polish_config.api_key if data.polish_config else None
    polish_base_url = request_polish_base_url
    enhance_model = data.enhance_config.model if data.enhance_config else None
    enhance_api_key = data.enhance_config.api_key if data.enhance_config else None
    enhance_base_url = request_enhance_base_url
    emotion_model = data.emotion_config.model if data.emotion_config else None
    emotion_api_key = data.emotion_config.api_key if data.emotion_config else None
    emotion_base_url = request_emotion_base_url

    if provider_config:
        polish_model = provider_config["polish_model"]
        polish_base_url = provider_config["base_url"]
        enhance_model = provider_config["enhance_model"]
        enhance_base_url = provider_config["base_url"]
        emotion_model = provider_config["emotion_model"]
        emotion_base_url = provider_config["base_url"] if provider_config["emotion_model"] else None
        if not data.polish_config:
            polish_api_key = None
            enhance_api_key = None
            emotion_api_key = None

    # 创建会话
    session_id = generate_session_id()
    session = OptimizationSession(
        user_id=user.id,
        session_id=session_id,
        original_text=data.original_text,
        processing_mode=data.processing_mode,
        billing_mode=data.billing_mode,
        credential_source=(
            "user_saved" if data.billing_mode == "byok" and not data.polish_config else "request"
            if data.billing_mode == "byok"
            else "system"
        ),
        charge_status="not_charged",
        charged_credits=0,
        current_stage=initial_stage,
        status="queued",
        progress=0.0,
        queued_at=utcnow(),
        polish_model=polish_model,
        polish_api_key=polish_api_key,
        polish_base_url=polish_base_url,
        enhance_model=enhance_model,
        enhance_api_key=enhance_api_key,
        enhance_base_url=enhance_base_url,
        emotion_model=emotion_model,
        emotion_api_key=emotion_api_key,
        emotion_base_url=emotion_base_url,
        project_id=project.id if project else None,
        task_title=data.task_title.strip() if data.task_title else None,
    )
    
    db.add(session)
    db.flush()
    if data.billing_mode == "platform" and data.processing_mode != "ai_detect_reduce":
        CreditService(db).hold_platform_credit(
            user,
            reason="optimization_start",
            session_id=session.id,
            amount=required_credits,
        )
        session.charge_status = "held"
        session.charged_credits = 0 if user.is_unlimited else required_credits
    elif data.billing_mode == "byok":
        session.charge_status = "not_charged"
        session.charged_credits = 0

    user.usage_count = usage_count + 1
    db.commit()
    db.refresh(session)
    if project:
        session.project = project
    
    if settings.INLINE_TASK_WORKER_ENABLED:
        background_tasks.add_task(run_optimization, session.id)
    
    return session


@router.post("/zhuque/browser/start", response_model=ZhuqueBrowserLaunchResponse)
async def start_zhuque_browser(
    sync_session: bool = True,
    mode: str = "remote_qr",
    user: User = Depends(get_current_user_with_legacy_fallback),
):
    """兼容旧路径：默认启动 VPS 可用的远程二维码登录会话。

    mode=local_window 时保留旧的服务端本机 Chrome 小窗同步能力，主要用于
    本地开发；公网/VPS 默认不要依赖服务端桌面。
    """
    if mode == "local_window":
        try:
            return _start_zhuque_wechat_capture(sync_session=sync_session, user=user)
        except TypeError:
            # Test/extension compatibility for older monkeypatches that only accept sync_session.
            return _start_zhuque_wechat_capture(sync_session=sync_session)
    payload = await zhuque_remote_login_service.start(user.id)
    getattr(zhuque_service, "reset_user", lambda _user_id: None)(user.id)
    return {
        "status": payload.get("status", "starting"),
        "auth_mode": payload.get("auth_mode", "headless_api"),
        "login_mode": payload.get("login_mode", "remote_wechat_qr"),
        "credential_file": payload.get("credential_file", ""),
        "sync_session": sync_session,
        "command": None,
        "message": payload.get("message", "请使用微信扫描二维码登录朱雀"),
        "session_id": payload.get("session_id", ""),
        "qr_image_data": payload.get("qr_image_data", ""),
        "expires_at": payload.get("expires_at", ""),
        "connected": bool(payload.get("connected")),
        "ready": bool(payload.get("ready")),
        "has_token": bool(payload.get("has_token")),
        "remaining_uses": payload.get("remaining_uses", -1),
        "user_name": payload.get("user_name", ""),
        "quota_text": payload.get("quota_text", ""),
    }


@router.get("/zhuque/browser/status", response_model=ZhuqueBrowserStatusResponse)
async def get_zhuque_browser_connection_status(
    user: User = Depends(get_current_user_with_legacy_fallback),
):
    """兼容旧路径：返回当前用户自己的朱雀凭证 / 无头 API 状态。"""
    return _get_zhuque_headless_status(user)


@router.get("/zhuque/browser/login-status", response_model=ZhuqueBrowserStatusResponse)
async def get_zhuque_remote_login_status(
    session_id: str = "",
    user: User = Depends(get_current_user_with_legacy_fallback),
):
    """轮询 VPS headless 朱雀二维码登录会话。"""
    payload = zhuque_remote_login_service.status(user.id, session_id or None)
    if payload.get("status") == "logged_in":
        getattr(zhuque_service, "reset_user", lambda _user_id: None)(user.id)
    return {
        "status": payload.get("status", "not_found"),
        "connected": bool(payload.get("connected")),
        "ready": bool(payload.get("ready")),
        "has_token": bool(payload.get("has_token")),
        "remaining_uses": payload.get("remaining_uses", -1),
        "button_enabled": payload.get("remaining_uses", -1) != 0,
        "auth_mode": payload.get("auth_mode", "headless_api"),
        "login_mode": payload.get("login_mode", "remote_wechat_qr"),
        "credential_file": payload.get("credential_file", ""),
        "user_name": payload.get("user_name", ""),
        "quota_text": payload.get("quota_text", ""),
        "captured_at": "",
        "message": payload.get("message", ""),
        "session_id": payload.get("session_id", ""),
        "qr_image_data": payload.get("qr_image_data", ""),
        "expires_at": payload.get("expires_at", ""),
    }


@router.post("/zhuque/browser/cancel", response_model=ZhuqueBrowserStatusResponse)
async def cancel_zhuque_remote_login(
    session_id: str = "",
    user: User = Depends(get_current_user_with_legacy_fallback),
):
    """取消当前用户的 VPS headless 朱雀扫码会话。"""
    payload = await zhuque_remote_login_service.cancel(user.id, session_id or None)
    return {
        "status": payload.get("status", "cancelled"),
        "connected": bool(payload.get("connected")),
        "ready": bool(payload.get("ready")),
        "has_token": bool(payload.get("has_token")),
        "remaining_uses": payload.get("remaining_uses", -1),
        "button_enabled": payload.get("remaining_uses", -1) != 0,
        "auth_mode": payload.get("auth_mode", "headless_api"),
        "login_mode": payload.get("login_mode", "remote_wechat_qr"),
        "credential_file": payload.get("credential_file", ""),
        "user_name": payload.get("user_name", ""),
        "quota_text": payload.get("quota_text", ""),
        "captured_at": "",
        "message": payload.get("message", ""),
        "session_id": payload.get("session_id", ""),
        "qr_image_data": payload.get("qr_image_data", ""),
        "expires_at": payload.get("expires_at", ""),
    }


@router.post("/zhuque/browser/logout", response_model=ZhuqueBrowserStatusResponse)
async def logout_zhuque_browser(
    user: User = Depends(get_current_user_with_legacy_fallback),
):
    """清除当前用户保存的朱雀登录凭证，后续检测回到未登录免费次数路径。"""
    payload = await zhuque_remote_login_service.logout(user.id)
    getattr(zhuque_service, "reset_user", lambda _user_id: None)(user.id)
    return {
        "status": payload.get("status", "logged_out"),
        "connected": False,
        "ready": False,
        "has_token": False,
        "remaining_uses": -1,
        "button_enabled": True,
        "auth_mode": payload.get("auth_mode", "headless_api"),
        "login_mode": payload.get("login_mode", "remote_wechat_qr"),
        "credential_file": payload.get("credential_file", ""),
        "user_name": "",
        "quota_text": "",
        "captured_at": "",
        "message": payload.get("message", "已退出朱雀登录，未登录时将使用朱雀免费次数"),
        "session_id": payload.get("session_id", ""),
        "qr_image_data": "",
        "expires_at": payload.get("expires_at", ""),
    }


@router.get("/zhuque/readiness", response_model=ZhuqueReadinessResponse)
async def get_zhuque_readiness(
    user: User = Depends(get_current_user_with_legacy_fallback),
):
    """读取当前用户朱雀页面是否可用于检测；不点击检测，不消耗朱雀次数。"""
    return await _zhuque_service_for_user(user).readiness()


@router.post("/zhuque/free-quota/refresh", response_model=ZhuqueReadinessResponse)
async def refresh_zhuque_free_quota(
    user: User = Depends(get_current_user_with_legacy_fallback),
):
    """主动探测当前用户朱雀剩余次数；不提交文本，不消耗朱雀检测次数。"""
    service = _zhuque_service_for_user(user)
    refresh_free_quota = getattr(service, "refresh_free_quota", None)
    if callable(refresh_free_quota):
        return await refresh_free_quota()
    return await service.readiness()


@router.post("/zhuque/preflight", response_model=ZhuqueReadinessResponse)
async def preflight_zhuque_task(
    payload: ZhuquePreflightRequest,
    user: User = Depends(get_current_user_with_legacy_fallback),
    db: Session = Depends(get_db),
):
    """带文本执行朱雀任务预检；不创建任务，不扣费。"""
    if payload.processing_mode != "ai_detect_reduce":
        return {
            "ready": True,
            "connected": True,
            "page_found": True,
            "has_token": False,
            "remaining_uses": -1,
            "button_enabled": True,
            "text_length": len(payload.original_text or ""),
            "text_length_ok": True,
            "estimated_first_round_credits": 0,
            "estimated_max_round_credits": 0,
            "message": "非朱雀模式无需预检",
            "actions": [],
        }

    if payload.billing_mode == "byok":
        ProviderConfigService(db).get_runtime_config(user)

    return _with_zhuque_cost_estimate(
        await _zhuque_service_for_user(user).readiness(payload.original_text),
        payload.original_text,
    )


@router.get("/status", response_model=QueueStatusResponse)
async def get_queue_status(
    session_id: str = None,
    user: User = Depends(get_current_user_with_legacy_fallback),
    db: Session = Depends(get_db)
):
    """获取队列状态"""
    status = await concurrency_manager.get_status(session_id)
    online_since = utcnow() - timedelta(seconds=ONLINE_USER_WINDOW_SECONDS)
    status["online_users"] = (
        db.query(User)
        .filter(
            User.is_active.is_(True),
            User.last_used.isnot(None),
            User.last_used >= online_since,
        )
        .count()
        or 0
    )
    return QueueStatusResponse(**status)


@router.get("/sessions", response_model=List[SessionResponse])
async def list_sessions(
    limit: int = 20,
    offset: int = 0,
    project_id: Optional[int] = None,
    user: User = Depends(get_current_user_with_legacy_fallback),
    db: Session = Depends(get_db)
):
    """列出用户的所有会话（支持分页）"""
    # 限制最大返回数量为100，避免一次性加载过多数据
    limit = min(limit, 100)
    
    # 查询会话及其原始文本长度和预览文本
    query = db.query(
        OptimizationSession,
        func.length(OptimizationSession.original_text).label('original_char_count'),
        func.substring(OptimizationSession.original_text, 1, 50).label('preview_text')
    ).options(
        joinedload(OptimizationSession.project),
        defer(OptimizationSession.original_text),
        defer(OptimizationSession.error_message)
    ).filter(
        OptimizationSession.user_id == user.id
    )

    if project_id is not None:
        if project_id == 0:
            query = query.filter(OptimizationSession.project_id.is_(None))
        else:
            query = query.filter(OptimizationSession.project_id == project_id)

    results = query.order_by(OptimizationSession.created_at.desc()).limit(limit).offset(offset).all()

    # 构造响应，手动注入 original_char_count 和 preview_text
    sessions = []
    for session, char_count, preview_text in results:
        session.original_char_count = char_count or 0
        session.preview_text = preview_text or ""
        sessions.append(session)
        
    return sessions


@router.get("/sessions/{session_id}", response_model=SessionDetailResponse)
async def get_session_detail(
    session_id: str,
    user: User = Depends(get_current_user_with_legacy_fallback),
    db: Session = Depends(get_db)
):
    """获取会话详情"""
    session = db.query(OptimizationSession).options(joinedload(OptimizationSession.project)).filter(
        OptimizationSession.session_id == session_id,
        OptimizationSession.user_id == user.id
    ).first()
    
    if not session:
        raise HTTPException(status_code=404, detail="会话不存在")
    
    # 获取段落
    segments = db.query(OptimizationSegment).filter(
        OptimizationSegment.session_id == session.id
    ).order_by(OptimizationSegment.segment_index).all()
    
    return SessionDetailResponse(
        **session.__dict__,
        project_title=session.project_title,
        segments=[seg.__dict__ for seg in segments]
    )


@router.patch("/sessions/{session_id}/project", response_model=SessionResponse)
async def update_session_project(
    session_id: str,
    payload: SessionProjectUpdateRequest,
    user: User = Depends(get_current_user_with_legacy_fallback),
    db: Session = Depends(get_db),
):
    """把单个会话归入指定论文项目；project_id=null 表示移回未归档。"""
    session = (
        db.query(OptimizationSession)
        .filter(
            OptimizationSession.session_id == session_id,
            OptimizationSession.user_id == user.id,
        )
        .first()
    )
    if not session:
        raise HTTPException(status_code=404, detail="会话不存在")

    project = None
    if payload.project_id is not None:
        project = (
            db.query(PaperProject)
            .filter(
                PaperProject.id == payload.project_id,
                PaperProject.user_id == user.id,
                PaperProject.is_archived.is_(False),
            )
            .first()
        )
        if not project:
            raise HTTPException(status_code=404, detail="论文项目不存在")

    session.project_id = project.id if project else None
    session.updated_at = utcnow()
    db.commit()
    db.refresh(session)
    if project:
        session.project = project
    else:
        session.project = None

    session.original_char_count = len(session.original_text or "")
    session.preview_text = (session.original_text or "")[:50]
    return session


@router.get("/sessions/{session_id}/progress", response_model=ProgressUpdate)
async def get_session_progress(
    session_id: str,
    user: User = Depends(get_current_user_with_legacy_fallback),
    db: Session = Depends(get_db)
):
    """获取会话进度"""
    # 查询完整会话对象，但避免急切加载关联对象
    session = (
        db.query(OptimizationSession)
        .options(joinedload(OptimizationSession.project))
        .filter(
            OptimizationSession.session_id == session_id,
            OptimizationSession.user_id == user.id,
        )
        .first()
    )
    
    if not session:
        raise HTTPException(status_code=404, detail="会话不存在")
    
    return ProgressUpdate(
        session_id=session.session_id,
        status=session.status,
        progress=session.progress,
        current_position=session.current_position,
        total_segments=session.total_segments,
        current_stage=session.current_stage,
        error_message=session.error_message
    )


@router.get("/sessions/{session_id}/stream")
async def stream_session_progress(
    session_id: str,
    request: Request,
    user: User = Depends(get_current_user_with_legacy_fallback),
    db: Session = Depends(get_db)
):
    """流式获取会话进度和内容"""
    # 验证用户权限
    session = db.query(OptimizationSession).filter(
        OptimizationSession.session_id == session_id,
        OptimizationSession.user_id == user.id
    ).first()
    
    if not session:
        raise HTTPException(status_code=404, detail="会话不存在")

    async def event_generator():
        queue = await stream_manager.connect(session_id)
        try:
            while True:
                if await request.is_disconnected():
                    break
                
                # 从队列获取消息，设置超时以便检查连接状态
                try:
                    message = await asyncio.wait_for(queue.get(), timeout=1.0)
                    yield message
                except asyncio.TimeoutError:
                    # 发送心跳注释以保持连接活跃
                    yield ": keep-alive\n\n"
                    
        finally:
            await stream_manager.disconnect(session_id, queue)

    return EventSourceResponse(event_generator())


@router.get("/sessions/{session_id}/changes", response_model=List[ChangeLogResponse])
async def get_session_changes(
    session_id: str,
    user: User = Depends(get_current_user_with_legacy_fallback),
    db: Session = Depends(get_db)
):
    """获取会话的变更对照"""
    session = db.query(OptimizationSession).filter(
        OptimizationSession.session_id == session_id,
        OptimizationSession.user_id == user.id
    ).first()
    
    if not session:
        raise HTTPException(status_code=404, detail="会话不存在")
    
    latest_log_subquery = db.query(
        ChangeLog.segment_index,
        ChangeLog.stage,
        func.max(ChangeLog.id).label("latest_id")
    ).filter(
        ChangeLog.session_id == session.id
    ).group_by(
        ChangeLog.segment_index,
        ChangeLog.stage
    ).subquery()

    change_logs = db.query(ChangeLog).join(
        latest_log_subquery,
        and_(
            ChangeLog.segment_index == latest_log_subquery.c.segment_index,
            ChangeLog.stage == latest_log_subquery.c.stage,
            ChangeLog.id == latest_log_subquery.c.latest_id
        )
    ).filter(
        ChangeLog.session_id == session.id
    ).order_by(
        ChangeLog.segment_index,
        case((ChangeLog.stage == "polish", 0), else_=1)
    ).all()

    parsed_changes = []
    for change in change_logs:
        detail = None
        if change.changes_detail:
            try:
                detail = json.loads(change.changes_detail)
            except json.JSONDecodeError:
                detail = {"raw": change.changes_detail}

        parsed_changes.append(
            ChangeLogResponse(
                id=change.id,
                segment_index=change.segment_index,
                stage=change.stage,
                before_text=change.before_text,
                after_text=change.after_text,
                changes_detail=detail,
                created_at=change.created_at
            )
        )

    return parsed_changes


@router.post("/sessions/{session_id}/export")
async def export_session(
    session_id: str,
    confirmation: ExportConfirmation,
    user: User = Depends(get_current_user_with_legacy_fallback),
    db: Session = Depends(get_db)
):
    """导出优化结果"""
    if not confirmation.acknowledge_academic_integrity:
        raise HTTPException(
            status_code=400,
            detail="必须确认学术诚信承诺"
        )
    
    session = db.query(OptimizationSession).filter(
        OptimizationSession.session_id == session_id,
        OptimizationSession.user_id == user.id
    ).first()
    
    if not session:
        raise HTTPException(status_code=404, detail="会话不存在")
    
    if session.status != "completed":
        raise HTTPException(status_code=400, detail="会话未完成")
    
    # 获取所有段落
    segments = db.query(OptimizationSegment).filter(
        OptimizationSegment.session_id == session.id
    ).order_by(OptimizationSegment.segment_index).all()

    if confirmation.export_format in {"aigc_report_docx", "aigc_report_md"}:
        if session.processing_mode != "ai_detect_reduce":
            raise HTTPException(status_code=400, detail="仅 AI检测+降重 会话支持导出 AIGC 检测报告")
        if not any(seg.zhuque_detect_result or seg.zhuque_detect_rate is not None for seg in segments):
            raise HTTPException(status_code=400, detail="当前会话没有可导出的朱雀 AIGC 检测结果")

        report_payload = _build_aigc_report_payload(session, segments)
        if confirmation.export_format == "aigc_report_md":
            return {
                "format": "aigc_report_md",
                "content": _build_aigc_report_markdown(report_payload),
                "filename": _build_aigc_report_filename(session, "md"),
                "mime_type": "text/markdown;charset=utf-8",
            }

        return {
            "format": "aigc_report_docx",
            "content_base64": _build_aigc_report_docx_base64(report_payload),
            "filename": _build_aigc_report_filename(session, "docx"),
            "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        }
    
    # 组合最终文本
    final_text = "\n\n".join([
        seg.zhuque_reduced_text or seg.enhanced_text or seg.polished_text or seg.original_text
        for seg in segments
    ])
    
    if confirmation.export_format == "md":
        return {
            "format": "md",
            "content": final_text,
            "filename": _build_export_filename(session, "md"),
            "mime_type": "text/markdown;charset=utf-8",
        }

    return {
        "format": "docx",
        "content_base64": _build_docx_base64(final_text),
        "filename": _build_export_filename(session, "docx"),
        "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }


@router.delete("/sessions/{session_id}")
async def delete_session(
    session_id: str,
    user: User = Depends(get_current_user_with_legacy_fallback),
    db: Session = Depends(get_db)
):
    """删除会话"""
    session = db.query(OptimizationSession).filter(
        OptimizationSession.session_id == session_id,
        OptimizationSession.user_id == user.id
    ).first()
    
    if not session:
        raise HTTPException(status_code=404, detail="会话不存在")
    
    db.delete(session)
    db.commit()
    
    return {"message": "会话已删除"}


@router.post("/sessions/{session_id}/retry")
async def retry_session(
    session_id: str,
    background_tasks: BackgroundTasks,
    data: Optional[SessionRetryRequest] = None,
    user: User = Depends(get_current_user_with_legacy_fallback),
    db: Session = Depends(get_db)
):
    """重新尝试处理失败的会话，继续未完成的段落"""
    session = db.query(OptimizationSession).filter(
        OptimizationSession.session_id == session_id,
        OptimizationSession.user_id == user.id
    ).first()

    if not session:
        raise HTTPException(status_code=404, detail="会话不存在")

    if session.status not in ["failed", "stopped"]:
        raise HTTPException(status_code=400, detail="仅可对失败或已停止的会话执行重试")

    requested_billing_mode = data.billing_mode if data else "keep"
    _apply_retry_billing_mode(
        session=session,
        user=user,
        requested_billing_mode=requested_billing_mode,
        db=db,
    )

    # 保留历史错误信息
    old_error = session.error_message or "未知错误"
    session.status = "queued"
    session.queued_at = utcnow()
    session.started_at = None
    session.finished_at = None
    session.worker_id = None
    session.error_message = f"[重试中] 上次失败原因: {old_error}"
    db.commit()

    if settings.INLINE_TASK_WORKER_ENABLED:
        background_tasks.add_task(run_optimization, session.id)

    return {
        "message": "已重新排队处理未完成段落",
        "billing_mode": session.billing_mode,
        "credential_source": session.credential_source,
    }


@router.post("/sessions/{session_id}/stop")
async def stop_session(
    session_id: str,
    user: User = Depends(get_current_user_with_legacy_fallback),
    db: Session = Depends(get_db)
):
    """停止正在进行中的会话"""
    session = db.query(OptimizationSession).filter(
        OptimizationSession.session_id == session_id,
        OptimizationSession.user_id == user.id
    ).first()

    if not session:
        raise HTTPException(status_code=404, detail="会话不存在")

    if session.status not in ["queued", "processing"]:
        raise HTTPException(status_code=400, detail="只能停止排队中或处理中的会话")

    # 更新状态为 stopped
    session.status = "stopped"
    session.error_message = "用户手动停止"
    db.commit()

    return {"message": "会话已停止"}
