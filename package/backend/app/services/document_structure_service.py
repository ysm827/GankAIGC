"""Document structure parsing and semantic segment classification.

This module owns the structure contract used by upload parsing and Zhuque
rewrite selection. It deliberately keeps MarkItDown/text-rule fallbacks so PDF
or DOCX structure extraction failure never turns into unsafe full-document
rewriting.
"""

from __future__ import annotations

import json
import os
import re
import tempfile
from html import unescape
from dataclasses import asdict, dataclass, field
from io import BytesIO
from typing import Any, Dict, Iterable, List, Optional, Tuple

from fastapi import HTTPException

from app.config import settings
from app.services.ai_service import count_text_length, split_text_into_segments
from app.services.mineru_service import MinerUError, MinerUParseResult, mineru_service


SEMANTIC_TYPE_TITLE = "TITLE"
SEMANTIC_TYPE_SECTION_HEADING = "SECTION_HEADING"
SEMANTIC_TYPE_ABSTRACT_HEADING = "ABSTRACT_HEADING"
SEMANTIC_TYPE_KEYWORDS_HEADING = "KEYWORDS_HEADING"
SEMANTIC_TYPE_ACK_HEADING = "ACK_HEADING"
SEMANTIC_TYPE_REFERENCE_HEADING = "REFERENCE_HEADING"
SEMANTIC_TYPE_TOC_HEADING = "TOC_HEADING"
SEMANTIC_TYPE_TOC_ITEM = "TOC_ITEM"
SEMANTIC_TYPE_ABSTRACT_BODY = "ABSTRACT_BODY"
SEMANTIC_TYPE_ACK_BODY = "ACK_BODY"
SEMANTIC_TYPE_BODY = "BODY"
SEMANTIC_TYPE_MIXED_HEADING_BODY = "MIXED_HEADING_BODY"
SEMANTIC_TYPE_KEYWORDS = "KEYWORDS"
SEMANTIC_TYPE_TABLE = "TABLE"
SEMANTIC_TYPE_CAPTION = "CAPTION"
SEMANTIC_TYPE_FORMULA = "FORMULA"
SEMANTIC_TYPE_REFERENCE_ITEM = "REFERENCE_ITEM"
SEMANTIC_TYPE_META = "META"
SEMANTIC_TYPE_HEADER_FOOTER = "HEADER_FOOTER"
SEMANTIC_TYPE_SHORT_TEXT = "SHORT_TEXT"
SEMANTIC_TYPE_UNKNOWN_PROTECTED = "UNKNOWN_PROTECTED"

SEMANTIC_SOURCE_TEXT_RULE = "text_rule"
SEMANTIC_SOURCE_DOCX_STYLE = "docx_style"
SEMANTIC_SOURCE_MINERU = "mineru"
SEMANTIC_SOURCE_MARKITDOWN_TEXT_RULE = "markitdown_text_rule"
SEMANTIC_SOURCE_MANUAL_TEXT_RULE = "manual_text_rule"
SEMANTIC_SOURCE_LEGACY_TEXT_RULE = "legacy_text_rule"

REDUCE_ALLOWED_TYPES = {
    SEMANTIC_TYPE_BODY,
    SEMANTIC_TYPE_MIXED_HEADING_BODY,
}

PROTECTED_SEMANTIC_TYPES = {
    SEMANTIC_TYPE_TITLE,
    SEMANTIC_TYPE_SECTION_HEADING,
    SEMANTIC_TYPE_ABSTRACT_HEADING,
    SEMANTIC_TYPE_KEYWORDS_HEADING,
    SEMANTIC_TYPE_ACK_HEADING,
    SEMANTIC_TYPE_REFERENCE_HEADING,
    SEMANTIC_TYPE_TOC_HEADING,
    SEMANTIC_TYPE_TOC_ITEM,
    SEMANTIC_TYPE_ABSTRACT_BODY,
    SEMANTIC_TYPE_ACK_BODY,
    SEMANTIC_TYPE_KEYWORDS,
    SEMANTIC_TYPE_TABLE,
    SEMANTIC_TYPE_CAPTION,
    SEMANTIC_TYPE_FORMULA,
    SEMANTIC_TYPE_REFERENCE_ITEM,
    SEMANTIC_TYPE_META,
    SEMANTIC_TYPE_HEADER_FOOTER,
    SEMANTIC_TYPE_SHORT_TEXT,
    SEMANTIC_TYPE_UNKNOWN_PROTECTED,
}


@dataclass
class SegmentSemanticDecision:
    semantic_type: str
    semantic_source: str
    semantic_confidence: float
    reduce_allowed: bool
    semantic_reason: str
    section: str = "BODY"
    length: int = 0

    def to_public_dict(self) -> Dict[str, Any]:
        return asdict(self)


@dataclass
class ParsedSegment:
    index: int
    text: str
    semantic_type: str
    semantic_source: str
    semantic_confidence: float
    reduce_allowed: bool
    semantic_reason: str
    char_start: int
    char_end: int
    page_number: Optional[int] = None
    bbox_json: Optional[str] = None

    def to_public_dict(self) -> Dict[str, Any]:
        return asdict(self)


@dataclass
class ParsedDocument:
    text: str
    segments: List[ParsedSegment]
    parser: str
    warnings: List[str] = field(default_factory=list)
    document_format: str = ""
    parse_engine: str = ""
    parse_fallback_used: bool = False
    parse_trace: Dict[str, Any] = field(default_factory=dict)

    @property
    def structure_summary(self) -> Dict[str, int]:
        counts: Dict[str, int] = {}
        for segment in self.segments:
            counts[segment.semantic_type] = counts.get(segment.semantic_type, 0) + 1
        return counts


class TextRuleSemanticClassifier:
    """Conservative semantic classifier for text-only and fallback paths."""

    def __init__(self, *, source: str = SEMANTIC_SOURCE_TEXT_RULE, skip_short_chars: Optional[int] = None):
        self.source = source
        self.skip_short_chars = max(
            0,
            int(skip_short_chars if skip_short_chars is not None else getattr(settings, "ZHUQUE_REDUCE_SKIP_SHORT_CHARS", 80) or 80),
        )

    def classify_segments(self, texts: Iterable[str]) -> List[SegmentSemanticDecision]:
        decisions: List[SegmentSemanticDecision] = []
        current_section = "BODY"
        reference_zone = False
        for text in texts:
            decision, current_section, reference_zone = self._classify_one(
                text,
                current_section=current_section,
                reference_zone=reference_zone,
            )
            decisions.append(decision)
        return decisions

    def classify_text(self, text: str) -> SegmentSemanticDecision:
        decision, _section, _reference_zone = self._classify_one(text, current_section="BODY", reference_zone=False)
        return decision

    def _classify_one(
        self,
        text: str,
        *,
        current_section: str,
        reference_zone: bool,
    ) -> Tuple[SegmentSemanticDecision, str, bool]:
        normalized = normalize_segment_line(text)
        classification_text = normalize_classification_text(normalized)
        length = count_text_length(classification_text)
        semantic_type = SEMANTIC_TYPE_BODY
        confidence = 0.75
        reason = "body_candidate"

        if is_reference_heading(classification_text):
            semantic_type = SEMANTIC_TYPE_REFERENCE_HEADING
            confidence = 0.98
            reason = "reference_heading"
            current_section = "REFERENCES"
            reference_zone = True
        elif is_toc_heading(classification_text):
            semantic_type = SEMANTIC_TYPE_TOC_HEADING
            confidence = 0.96
            reason = "toc_heading"
            current_section = "TOC"
            reference_zone = False
        elif reference_zone and is_reference_item(classification_text):
            semantic_type = SEMANTIC_TYPE_REFERENCE_ITEM
            confidence = 0.9
            reason = "reference_zone_item"
        elif reference_zone and is_repeating_pdf_artifact(classification_text):
            semantic_type = SEMANTIC_TYPE_HEADER_FOOTER
            confidence = 0.88
            reason = "reference_zone_repeating_artifact"
        elif reference_zone:
            semantic_type = SEMANTIC_TYPE_REFERENCE_ITEM
            confidence = 0.72
            reason = "reference_zone_continuation"
        elif is_standalone_reference_item(classification_text):
            semantic_type = SEMANTIC_TYPE_REFERENCE_ITEM
            confidence = 0.86
            reason = "standalone_reference_item"
        elif current_section == "TOC" and is_toc_item(classification_text):
            semantic_type = SEMANTIC_TYPE_TOC_ITEM
            confidence = 0.9
            reason = "toc_item"
        elif is_abstract_heading(classification_text):
            semantic_type = SEMANTIC_TYPE_ABSTRACT_HEADING
            confidence = 0.98
            reason = "abstract_heading"
            current_section = "ABSTRACT"
            reference_zone = False
        elif is_ack_heading(classification_text):
            semantic_type = SEMANTIC_TYPE_ACK_HEADING
            confidence = 0.98
            reason = "ack_heading"
            current_section = "ACK"
            reference_zone = False
        elif is_keywords_line(classification_text):
            semantic_type = SEMANTIC_TYPE_KEYWORDS
            confidence = 0.92
            reason = "keywords_line"
            reference_zone = False
        elif is_mixed_heading_body(classification_text):
            semantic_type = SEMANTIC_TYPE_MIXED_HEADING_BODY
            confidence = 0.78
            reason = "mixed_heading_body"
            current_section = "BODY"
            reference_zone = False
        elif is_section_heading(classification_text):
            semantic_type = SEMANTIC_TYPE_SECTION_HEADING
            confidence = 0.86
            reason = "section_heading"
            current_section = "BODY"
            reference_zone = False
        elif is_caption(classification_text):
            semantic_type = SEMANTIC_TYPE_CAPTION
            confidence = 0.82
            reason = "caption"
            reference_zone = False
        elif is_formula_or_metric(classification_text):
            semantic_type = SEMANTIC_TYPE_FORMULA
            confidence = 0.84
            reason = "formula_or_metric"
            reference_zone = False
        elif current_section == "ABSTRACT":
            semantic_type = SEMANTIC_TYPE_ABSTRACT_BODY
            confidence = 0.86
            reason = "abstract_body_protected"
        elif current_section == "ACK":
            semantic_type = SEMANTIC_TYPE_ACK_BODY
            confidence = 0.86
            reason = "ack_body_protected"
        elif is_meta_line(classification_text):
            semantic_type = SEMANTIC_TYPE_META
            confidence = 0.78
            reason = "metadata_line"
            reference_zone = False
        elif length < self.skip_short_chars:
            semantic_type = SEMANTIC_TYPE_SHORT_TEXT
            confidence = 0.7
            reason = "short_text"
            reference_zone = False

        return (
            SegmentSemanticDecision(
                semantic_type=semantic_type,
                semantic_source=self.source,
                semantic_confidence=confidence,
                reduce_allowed=semantic_type in REDUCE_ALLOWED_TYPES,
                semantic_reason=reason,
                section=current_section,
                length=length,
            ),
            current_section,
            reference_zone,
        )



class DocumentStructureService:
    def parse_uploaded_document(self, content: bytes, extension: str, filename: str = "") -> ParsedDocument:
        extension = extension.lower()
        if extension in {".md", ".markdown"}:
            text, warnings = decode_markdown_upload(content)
            normalized = normalize_parsed_document_text(text)
            return build_parsed_document_from_text(
                normalized,
                parser="markdown",
                document_format=extension.lstrip("."),
                semantic_source=SEMANTIC_SOURCE_MANUAL_TEXT_RULE,
                warnings=warnings,
                parse_engine="markdown",
                trace={"engine": "markdown"},
            )
        if extension == ".docx":
            return self._parse_docx(content)
        if extension == ".pdf":
            return self._parse_pdf(content, filename or "paper.pdf")
        raise HTTPException(status_code=400, detail="仅支持上传 Word(.docx)、PDF(.pdf) 和 Markdown(.md/.markdown) 文件")

    def classify_manual_text(self, text: str) -> ParsedDocument:
        normalized = normalize_parsed_document_text(text)
        return build_parsed_document_from_text(
            normalized,
            parser="manual_text_rule",
            document_format="text",
            semantic_source=SEMANTIC_SOURCE_MANUAL_TEXT_RULE,
            warnings=[],
            parse_engine="manual_text_rule",
            trace={"engine": "manual_text_rule"},
        )

    def _parse_docx(self, content: bytes) -> ParsedDocument:
        try:
            from docx import Document as DocxDocument

            document = DocxDocument(BytesIO(content))
            raw_segments: List[Tuple[str, Optional[SegmentSemanticDecision], Optional[int], Optional[str]]] = []
            for paragraph in document.paragraphs:
                text = normalize_segment_line(paragraph.text)
                if not text:
                    continue
                style_name = getattr(getattr(paragraph, "style", None), "name", "") or ""
                style_decision = classify_docx_style(style_name)
                raw_segments.append((text, style_decision, None, None))

            for section in document.sections:
                for container in (section.header, section.footer):
                    for paragraph in container.paragraphs:
                        text = normalize_segment_line(paragraph.text)
                        if not text:
                            continue
                        raw_segments.append((
                            text,
                            SegmentSemanticDecision(
                                semantic_type=SEMANTIC_TYPE_HEADER_FOOTER,
                                semantic_source=SEMANTIC_SOURCE_DOCX_STYLE,
                                semantic_confidence=0.9,
                                reduce_allowed=False,
                                semantic_reason="docx_header_footer",
                                section="HEADER_FOOTER",
                                length=count_text_length(text),
                            ),
                            None,
                            None,
                        ))

            if raw_segments:
                parsed = build_parsed_document_from_raw_segments(
                    raw_segments,
                    parser="python_docx",
                    document_format="docx",
                    fallback_source=SEMANTIC_SOURCE_DOCX_STYLE,
                    warnings=[],
                    parse_engine="python_docx",
                    trace={"engine": "python_docx", "style_segments": sum(1 for _text, decision, _page, _bbox in raw_segments if decision)},
                )
                if parsed.text.strip():
                    return parsed
        except Exception as exc:
            fallback = parse_document_with_markitdown(content, ".docx")
            fallback.warnings.append(f"Word 样式解析失败，已回退 MarkItDown：{exc}")
            return fallback

        fallback = parse_document_with_markitdown(content, ".docx")
        fallback.warnings.append("Word 样式解析未提取到有效文本，已回退 MarkItDown")
        return fallback

    def _parse_pdf(self, content: bytes, filename: str) -> ParsedDocument:
        engine = (settings.PDF_STRUCTURE_ENGINE or "mineru").strip().lower()
        if engine == "markitdown":
            return parse_document_with_markitdown(content, ".pdf")
        if engine != "mineru":
            fallback = parse_document_with_markitdown(content, ".pdf")
            fallback.warnings.append(f"PDF 结构解析引擎 {engine} 不受支持，已使用 MarkItDown")
            fallback.parse_trace = {**fallback.parse_trace, "requested_engine": engine, "fallback_reason": "unsupported_engine"}
            return fallback

        try:
            mineru_result = mineru_service.parse_pdf(filename, content)
            parsed = build_parsed_document_from_mineru_result(mineru_result)
            if parsed.text.strip():
                return parsed
            raise MinerUError("MinerU 未提取到有效文本")
        except Exception as exc:
            fallback = parse_document_with_markitdown(content, ".pdf")
            fallback.parse_fallback_used = True
            fallback.warnings.append(f"MinerU 解析失败，已回退 MarkItDown：{exc}")
            fallback.parse_trace = {
                **fallback.parse_trace,
                "fallback_from": "mineru",
                "fallback_reason": type(exc).__name__,
                "fallback_message": str(exc)[:200],
            }
            return fallback


# Public singleton; callers can monkeypatch module functions or this object in tests.
document_structure_service = DocumentStructureService()


def normalize_segment_line(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "").strip())


def normalize_classification_text(text: str) -> str:
    normalized = normalize_segment_line(text)
    normalized = re.sub(r"^#{1,6}\s*", "", normalized).strip()
    normalized = re.sub(r"^(?:>\s*)+", "", normalized).strip()
    normalized = re.sub(r"^\*\*(.+?)\*\*", r"\1", normalized).strip()
    normalized = re.sub(r"^__(.+?)__", r"\1", normalized).strip()
    for marker in ("**", "__", "*", "_"):
        if normalized.startswith(marker) and normalized.endswith(marker) and len(normalized) > 2 * len(marker):
            normalized = normalized[len(marker):-len(marker)].strip()
    return normalized


def is_reference_heading(text: str) -> bool:
    normalized = normalize_classification_text(text).strip("：: ")
    compact = re.sub(r"\s+", "", normalized).lower()
    if compact in {"参考文献", "参考资料", "主要参考文献"} or normalized.lower() in {
        "references",
        "reference",
        "bibliography",
        "bibliographical references",
        "参考书目",
        "works cited",
    }:
        return True
    heading_pattern = (
        r"^(?:"
        r"\d+(?:\.\d+)*[、.]?|"
        r"[一二三四五六七八九十]+[、.]|"
        r"第[一二三四五六七八九十\d]+[章节]|"
        r"[ⅠⅡⅢⅣⅤⅥⅦⅧⅨⅩ]+|"
        r"[IVXLCDM]+[.]?"
        r")\s*(?:参考文献|参考资料|主要参考文献|references?|bibliography|works cited)$"
    )
    return bool(re.match(heading_pattern, normalized, re.IGNORECASE))


def is_toc_heading(text: str) -> bool:
    normalized = normalize_classification_text(text).strip("：: ")
    compact = re.sub(r"\s+", "", normalized).lower()
    return compact == "目录" or normalized.lower() in {"contents", "table of contents"}


def is_toc_item(text: str) -> bool:
    normalized = normalize_classification_text(text)
    if not normalized:
        return False
    if re.match(r"^\[[^\]]+\]\(#_?toc", normalized, re.IGNORECASE):
        return True
    if re.match(r"^\d+(?:\.\d+)*\s+.{1,100}\s+\d+$", normalized):
        return True
    if re.match(r"^.{1,100}(?:\.{2,}|…{2,}|·{2,}|\s{2,})\s*(?:\d+|[ivxlcdm]+|[IVXLCDM]+)$", normalized):
        return True
    return False


def is_abstract_heading(text: str) -> bool:
    normalized = normalize_classification_text(text).strip("：: ")
    return normalized.lower() in {"摘要", "abstract"}


def is_ack_heading(text: str) -> bool:
    normalized = normalize_classification_text(text).strip("：: ")
    compact = re.sub(r"\s+", "", normalized).lower()
    return compact in {"致谢", "谢辞", "鸣谢"} or normalized.lower() in {
        "acknowledgement",
        "acknowledgements",
        "acknowledgment",
        "acknowledgments",
        "thanks",
    }


def is_keywords_line(text: str) -> bool:
    normalized = normalize_classification_text(text)
    return bool(re.match(r"^(关键词|关键字|keywords?)\s*[:：]", normalized, re.IGNORECASE))


def is_section_heading(text: str) -> bool:
    normalized = normalize_classification_text(text).strip()
    length = count_text_length(normalized)
    if not normalized or length > 80:
        return False
    lower = normalized.lower().strip("：: ")
    section_words = {
        "引言", "绪论", "前言", "研究背景", "研究方法", "方法", "材料与方法",
        "实验", "实验结果", "结果", "讨论", "结论", "总结", "展望",
        "introduction", "background", "methods", "materials and methods",
        "results", "discussion", "conclusion", "conclusions", "limitations",
    }
    if lower in section_words:
        return True
    return bool(
        re.match(r"^((第[一二三四五六七八九十\d]+[章节])|([一二三四五六七八九十]+[、.])|(\d+(\.\d+)*[、.]?))\s*[\u4e00-\u9fffA-Za-z].{0,60}$", normalized)
        or re.match(r"^(chapter|section)\s+([\divxlcdm]+|one|two|three|four|five|six|seven|eight|nine|ten)\b.{0,60}$", normalized, re.IGNORECASE)
        or re.match(r"^[ⅠⅡⅢⅣⅤⅥⅦⅧⅨⅩ]+\s*[\u4e00-\u9fffA-Za-z].{0,60}$", normalized)
    )


def is_mixed_heading_body(text: str) -> bool:
    normalized = normalize_classification_text(text).strip()
    length = count_text_length(normalized)
    if length < 80:
        return False
    return bool(
        re.match(r"^((第[一二三四五六七八九十\d]+[章节])|(\d+(\.\d+)*[、.]?))\s*[\u4e00-\u9fffA-Za-z].{4,}", normalized)
    )


def is_caption(text: str) -> bool:
    normalized = normalize_classification_text(text)
    return bool(re.match(r"^(图|表)\s*\d+|^(figure|fig\.|table)\s*\d+", normalized, re.IGNORECASE))


def is_reference_item(text: str) -> bool:
    normalized = normalize_classification_text(text)
    if not normalized:
        return False
    if re.match(r"^(\[\d+\]|\d+[.．、])", normalized):
        if re.search(r"\b\d{4}\b|\[J\]|\[D\]|\[M\]|\[N\]|doi|https?://", normalized, re.IGNORECASE):
            return True
    if re.search(r"\bdoi\s*[:：]|https?://|www\.", normalized, re.IGNORECASE):
        return True
    if re.search(r"\(\d{4}[a-z]?\)|\b\d{4}\b", normalized) and re.search(r"\bet al\.|[A-Z][a-z]+,\s*[A-Z]\.", normalized):
        return True
    return False


def is_standalone_reference_item(text: str) -> bool:
    normalized = normalize_classification_text(text)
    if not normalized:
        return False
    if has_reference_item_prefix(normalized):
        return bool(re.search(r"\b\d{4}\b|\[J\]|\[D\]|\[M\]|\[N\]|doi|https?://", normalized, re.IGNORECASE))
    return bool(
        re.search(r"\(\d{4}[a-z]?\)|\b\d{4}\b", normalized)
        and re.search(r"\bet al\.|[A-Z][a-z]+,\s*[A-Z]\.", normalized)
    )


def is_formula_or_metric(text: str) -> bool:
    normalized = normalize_classification_text(text)
    if not normalized:
        return False
    natural_chars = len(re.findall(r"[\u4e00-\u9fffA-Za-z]", normalized))
    symbol_chars = len(re.findall(r"[=<>±%＋+\-*/^_√∑∫≈≤≥]", normalized))
    digit_chars = len(re.findall(r"\d", normalized))
    total_chars = max(len(normalized), 1)
    if symbol_chars >= 2 and (symbol_chars + digit_chars) / total_chars >= 0.35:
        return True
    metric_pattern = r"\b(acc|accuracy|f1|auc|rmse|mae|mse|p\s*[<=>]|r2|dice|iou)\b"
    if re.search(metric_pattern, normalized, re.IGNORECASE) and (digit_chars + symbol_chars) / total_chars >= 0.25:
        return True
    return natural_chars <= 6 and (digit_chars + symbol_chars) >= 4


def is_meta_line(text: str) -> bool:
    normalized = normalize_classification_text(text)
    if not normalized:
        return False
    if re.search(r"@|邮箱|通讯作者|基金项目|作者简介|单位[:：]|学院|大学|实验室|原创性声明|版权使用授权书|学位论文", normalized, re.IGNORECASE):
        return count_text_length(normalized) < 160
    return False


def is_repeating_pdf_artifact(text: str) -> bool:
    normalized = normalize_classification_text(text).strip()
    if not normalized:
        return False
    if is_isolated_page_number(normalized):
        return True
    if has_reference_item_prefix(normalized):
        return False
    if re.match(r"^doi\s*[:：]\s*\S+", normalized, re.IGNORECASE):
        return count_text_length(normalized) <= 160
    if re.search(r"\bISSN\b|\bVolume\b|\bIssue\b|Journal of", normalized, re.IGNORECASE):
        return count_text_length(normalized) <= 160
    return False


def is_isolated_page_number(text: str) -> bool:
    normalized = normalize_classification_text(text).strip()
    return bool(re.match(r"^(?:\d{1,4}|[ivxlcdm]{1,8}|[IVXLCDM]{1,8})$", normalized))


def has_reference_item_prefix(text: str) -> bool:
    normalized = normalize_classification_text(text)
    return bool(re.match(r"^(\[\d+\]|\d+[.．、])", normalized))


def classify_docx_style(style_name: str) -> Optional[SegmentSemanticDecision]:
    normalized = (style_name or "").strip().lower()
    if not normalized:
        return None
    length = 0
    if normalized == "title" or normalized in {"标题"}:
        return SegmentSemanticDecision(SEMANTIC_TYPE_TITLE, SEMANTIC_SOURCE_DOCX_STYLE, 0.95, False, "docx_title_style", length=length)
    if re.match(r"^(heading|标题)\s*[1-4]$", normalized):
        return SegmentSemanticDecision(SEMANTIC_TYPE_SECTION_HEADING, SEMANTIC_SOURCE_DOCX_STYLE, 0.94, False, "docx_heading_style", length=length)
    if re.match(r"^toc\s*[1-9]$", normalized) or re.match(r"^目录\s*[1-9]$", normalized):
        return SegmentSemanticDecision(SEMANTIC_TYPE_TOC_ITEM, SEMANTIC_SOURCE_DOCX_STYLE, 0.94, False, "docx_toc_style", section="TOC", length=length)
    return None


def merge_pdf_line_segments(
    raw_segments: List[Tuple[str, Optional[SegmentSemanticDecision], Optional[int], Optional[str]]],
) -> List[Tuple[str, Optional[SegmentSemanticDecision], Optional[int], Optional[str]]]:
    """Merge PDF extractor text-line items into paragraph-like body segments."""
    merged: List[Tuple[str, Optional[SegmentSemanticDecision], Optional[int], Optional[str]]] = []
    buffer: List[str] = []
    buffer_page: Optional[int] = None

    def flush_buffer() -> None:
        nonlocal buffer, buffer_page
        if not buffer:
            return
        text = join_pdf_body_lines(buffer)
        if text:
            merged.append((text, None, buffer_page, None))
        buffer = []
        buffer_page = None

    for text, decision, page_no, bbox_json in raw_segments:
        if is_pdf_body_line_candidate(text, decision):
            if buffer and buffer_page != page_no:
                flush_buffer()
            buffer.append(text)
            buffer_page = page_no
            continue

        flush_buffer()
        merged.append((text, decision, page_no, bbox_json))

    flush_buffer()
    return merged


def is_pdf_body_line_candidate(text: str, decision: Optional[SegmentSemanticDecision]) -> bool:
    if not text:
        return False
    if decision and decision.semantic_type not in REDUCE_ALLOWED_TYPES:
        return False
    normalized = normalize_classification_text(text)
    if is_reference_heading(normalized) or is_toc_heading(normalized) or is_abstract_heading(normalized) or is_ack_heading(normalized):
        return False
    if is_section_heading(normalized) or is_caption(normalized) or is_formula_or_metric(normalized):
        return False
    return True


def join_pdf_body_lines(lines: List[str]) -> str:
    joined = ""
    for line in lines:
        current = normalize_segment_line(line)
        if not current:
            continue
        if not joined:
            joined = current
            continue
        if should_join_pdf_line_without_space(joined, current):
            joined = f"{joined}{current}"
        else:
            joined = f"{joined} {current}"
    return normalize_segment_line(joined)


def should_join_pdf_line_without_space(previous: str, current: str) -> bool:
    if not previous or not current:
        return False
    return bool(re.search(r"[\u4e00-\u9fff]$", previous) and re.match(r"^[\u4e00-\u9fff]", current))


def build_parsed_document_from_mineru_result(result: MinerUParseResult) -> ParsedDocument:
    raw_segments = cleanup_pdf_repeating_artifacts(
        mineru_content_items_to_raw_segments(result.content_items)
    )
    parsed = build_parsed_document_from_raw_segments(
        raw_segments,
        parser="mineru",
        document_format="pdf",
        fallback_source=SEMANTIC_SOURCE_MINERU,
        warnings=[],
        parse_engine="mineru",
        trace={
            "engine": "mineru",
            "api": "v4",
            "batch_id": result.batch_id,
            "trace_id": result.trace_id,
            "model_version": result.model_version,
            "content_item_count": len(result.content_items),
            "content_list_name": result.content_list_name,
        },
    )
    return parsed


def mineru_content_items_to_raw_segments(
    items: List[Dict[str, Any]],
) -> List[Tuple[str, Optional[SegmentSemanticDecision], Optional[int], Optional[str]]]:
    raw_segments: List[Tuple[str, Optional[SegmentSemanticDecision], Optional[int], Optional[str]]] = []
    for item in items:
        item_type = str(item.get("type") or "").strip()
        page_number = mineru_page_number(item)
        bbox_json = mineru_bbox_json(item)

        if item_type == "text":
            text = normalize_segment_line(mineru_scalar_text(item, "text"))
            if not text:
                continue
            text_level = mineru_text_level(item)
            decision = None
            if text_level > 0:
                decision = mineru_semantic_decision(
                    SEMANTIC_TYPE_SECTION_HEADING,
                    text,
                    "mineru_text_level_heading",
                    confidence=0.94,
                )
            raw_segments.append((text, decision, page_number, bbox_json))
            continue

        if item_type == "table":
            text = mineru_join_text_parts(
                mineru_text_list(item, "table_caption")
                + [strip_html_text(mineru_scalar_text(item, "table_body"))]
                + mineru_text_list(item, "table_footnote")
            )
            if text:
                raw_segments.append((
                    text,
                    mineru_semantic_decision(SEMANTIC_TYPE_TABLE, text, "mineru_table", confidence=0.96, section="TABLE"),
                    page_number,
                    bbox_json,
                ))
            continue

        if item_type == "equation":
            text = normalize_segment_line(mineru_scalar_text(item, "text"))
            if text:
                raw_segments.append((
                    text,
                    mineru_semantic_decision(SEMANTIC_TYPE_FORMULA, text, "mineru_equation", confidence=0.96, section="FORMULA"),
                    page_number,
                    bbox_json,
                ))
            continue

        if item_type == "image":
            text = mineru_join_text_parts(
                mineru_text_list(item, "image_caption")
                + mineru_text_list(item, "image_footnote")
            )
            if text:
                raw_segments.append((
                    text,
                    mineru_semantic_decision(SEMANTIC_TYPE_CAPTION, text, "mineru_image_caption", confidence=0.9, section="CAPTION"),
                    page_number,
                    bbox_json,
                ))
            continue

        if item_type == "chart":
            text = mineru_join_text_parts(
                mineru_text_list(item, "chart_caption")
                + [mineru_scalar_text(item, "content")]
                + mineru_text_list(item, "chart_footnote")
            )
            if text:
                raw_segments.append((
                    text,
                    mineru_semantic_decision(SEMANTIC_TYPE_CAPTION, text, "mineru_chart_caption", confidence=0.9, section="CAPTION"),
                    page_number,
                    bbox_json,
                ))
            continue

        if item_type == "list":
            sub_type = str(item.get("sub_type") or "").strip()
            for list_item_text in mineru_text_list(item, "list_items"):
                text = normalize_segment_line(list_item_text)
                if not text:
                    continue
                decision = None
                if sub_type == "ref_text":
                    decision = mineru_semantic_decision(
                        SEMANTIC_TYPE_REFERENCE_ITEM,
                        text,
                        "mineru_reference_list_item",
                        confidence=0.94,
                        section="REFERENCES",
                    )
                raw_segments.append((text, decision, page_number, bbox_json))
            continue

        if item_type == "ref_text":
            text = normalize_segment_line(mineru_scalar_text(item, "text"))
            if text:
                raw_segments.append((
                    text,
                    mineru_semantic_decision(
                        SEMANTIC_TYPE_REFERENCE_ITEM,
                        text,
                        "mineru_reference_text",
                        confidence=0.94,
                        section="REFERENCES",
                    ),
                    page_number,
                    bbox_json,
                ))
            continue

        if item_type == "code":
            text = mineru_join_text_parts(
                mineru_text_list(item, "code_caption")
                + [mineru_scalar_text(item, "code_body")]
                + mineru_text_list(item, "code_footnote")
            )
            if text:
                raw_segments.append((
                    text,
                    mineru_semantic_decision(SEMANTIC_TYPE_META, text, "mineru_code", confidence=0.88, section="META"),
                    page_number,
                    bbox_json,
                ))
            continue

        if item_type in {"header", "footer", "page_number"}:
            text = normalize_segment_line(mineru_scalar_text(item, "text"))
            if text:
                raw_segments.append((
                    text,
                    mineru_semantic_decision(
                        SEMANTIC_TYPE_HEADER_FOOTER,
                        text,
                        f"mineru_{item_type}",
                        confidence=0.9,
                        section="HEADER_FOOTER",
                    ),
                    page_number,
                    bbox_json,
                ))
            continue

        if item_type in {"aside_text", "page_footnote"}:
            text = normalize_segment_line(mineru_scalar_text(item, "text"))
            if text:
                raw_segments.append((
                    text,
                    mineru_semantic_decision(SEMANTIC_TYPE_META, text, f"mineru_{item_type}", confidence=0.86, section="META"),
                    page_number,
                    bbox_json,
                ))
            continue

    return raw_segments


def cleanup_pdf_repeating_artifacts(
    raw_segments: List[Tuple[str, Optional[SegmentSemanticDecision], Optional[int], Optional[str]]],
) -> List[Tuple[str, Optional[SegmentSemanticDecision], Optional[int], Optional[str]]]:
    """Drop repeated PDF headers/footers and isolated page-number artifacts.

    MinerU can expose visual headers/footers as ordinary text when the source
    PDF embeds them in the text layer. Keeping every repeated journal header
    inside References breaks the reference-zone classifier and pollutes the
    textarea. We remove only high-confidence artifacts:
    - isolated page numbers;
    - same normalized short header/footer-ish line repeated on 2+ pages.
    """
    normalized_counts: Dict[str, int] = {}
    page_sets: Dict[str, set[int]] = {}
    for text, _decision, page_number, _bbox_json in raw_segments:
        normalized = normalize_pdf_artifact_key(text)
        if not normalized:
            continue
        normalized_counts[normalized] = normalized_counts.get(normalized, 0) + 1
        if page_number is not None:
            page_sets.setdefault(normalized, set()).add(page_number)

    cleaned: List[Tuple[str, Optional[SegmentSemanticDecision], Optional[int], Optional[str]]] = []
    for text, decision, page_number, bbox_json in raw_segments:
        normalized = normalize_pdf_artifact_key(text)
        if decision and decision.semantic_type == SEMANTIC_TYPE_HEADER_FOOTER:
            continue
        if normalized and is_isolated_page_number(text):
            continue
        if normalized and has_reference_item_prefix(text):
            cleaned.append((text, decision, page_number, bbox_json))
            continue
        repeated_on_pages = len(page_sets.get(normalized, set())) >= 2
        repeated_count = normalized_counts.get(normalized, 0) >= 2
        if normalized and (repeated_on_pages or repeated_count) and is_repeating_pdf_artifact(text):
            continue
        cleaned.append((text, decision, page_number, bbox_json))
    return cleaned


def normalize_pdf_artifact_key(text: str) -> str:
    normalized = normalize_classification_text(text)
    normalized = re.sub(r"\s+", " ", normalized).strip().lower()
    return normalized


def mineru_semantic_decision(
    semantic_type: str,
    text: str,
    reason: str,
    *,
    confidence: float,
    section: str = "BODY",
) -> SegmentSemanticDecision:
    return SegmentSemanticDecision(
        semantic_type=semantic_type,
        semantic_source=SEMANTIC_SOURCE_MINERU,
        semantic_confidence=confidence,
        reduce_allowed=semantic_type in REDUCE_ALLOWED_TYPES,
        semantic_reason=reason,
        section=section,
        length=count_text_length(text),
    )


def mineru_scalar_text(item: Dict[str, Any], field_name: str) -> str:
    value = item.get(field_name)
    return value if isinstance(value, str) else ""


def mineru_text_list(item: Dict[str, Any], field_name: str) -> List[str]:
    value = item.get(field_name)
    if isinstance(value, str):
        return [value]
    if isinstance(value, list):
        return [entry for entry in value if isinstance(entry, str)]
    return []


def mineru_join_text_parts(parts: List[str]) -> str:
    normalized_parts = [normalize_segment_line(part) for part in parts if normalize_segment_line(part)]
    return normalize_segment_line(" ".join(normalized_parts))


def strip_html_text(text: str) -> str:
    if not text:
        return ""
    without_tags = re.sub(r"<[^>]+>", " ", text)
    return normalize_segment_line(unescape(without_tags))


def mineru_text_level(item: Dict[str, Any]) -> int:
    value = item.get("text_level", 0)
    if isinstance(value, bool):
        return 0
    if isinstance(value, int):
        return max(0, value)
    return 0


def mineru_page_number(item: Dict[str, Any]) -> Optional[int]:
    value = item.get("page_idx")
    if isinstance(value, bool):
        return None
    if isinstance(value, int) and value >= 0:
        return value + 1
    return None


def mineru_bbox_json(item: Dict[str, Any]) -> str:
    metadata: Dict[str, Any] = {"mineru_type": str(item.get("type") or "")}
    if "bbox" in item:
        metadata["bbox"] = item["bbox"]
    if "text_level" in item:
        metadata["text_level"] = item["text_level"]
    if "sub_type" in item:
        metadata["sub_type"] = item["sub_type"]
    return json.dumps(metadata, ensure_ascii=False, separators=(",", ":"))


def decode_markdown_upload(content: bytes) -> Tuple[str, List[str]]:
    warnings: List[str] = []
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return content.decode(encoding), warnings
        except UnicodeDecodeError:
            continue
    warnings.append("Markdown 文件编码无法准确识别，已使用替换模式读取")
    return content.decode("utf-8", errors="replace"), warnings


def parse_document_with_markitdown(content: bytes, extension: str) -> ParsedDocument:
    try:
        from markitdown import MarkItDown
    except ImportError as exc:
        raise HTTPException(status_code=500, detail="服务器未安装 MarkItDown，暂时无法解析 Word/PDF 文件") from exc

    temp_path = None
    extension = extension if extension in {".docx", ".pdf"} else ".docx"
    try:
        md = MarkItDown(enable_plugins=False)
        convert_stream = getattr(md, "convert_stream", None)
        if callable(convert_stream):
            result = convert_stream(BytesIO(content), file_extension=extension)
        else:
            with tempfile.NamedTemporaryFile(suffix=extension, delete=False) as temp_file:
                temp_file.write(content)
                temp_path = temp_file.name
            convert_local = getattr(md, "convert_local", None)
            result = convert_local(temp_path) if callable(convert_local) else md.convert(temp_path)
    except HTTPException:
        raise
    except Exception as exc:
        file_label = "PDF" if extension == ".pdf" else "Word"
        raise HTTPException(status_code=400, detail=f"{file_label} 文件解析失败: {exc}") from exc
    finally:
        if temp_path:
            try:
                os.unlink(temp_path)
            except OSError:
                pass

    text = getattr(result, "text_content", "") or ""
    normalized = normalize_parsed_document_text(text)
    return build_parsed_document_from_text(
        normalized,
        parser="markitdown",
        document_format=extension.lstrip("."),
        semantic_source=SEMANTIC_SOURCE_MARKITDOWN_TEXT_RULE,
        warnings=[],
        parse_engine="markitdown",
        trace={"engine": "markitdown"},
    )


def normalize_parsed_document_text(text: str) -> str:
    normalized = (text or "").replace("\r\n", "\n").replace("\r", "\n")
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    return normalized.strip()


def build_parsed_document_from_text(
    text: str,
    *,
    parser: str,
    document_format: str,
    semantic_source: str,
    warnings: List[str],
    parse_engine: str,
    trace: Dict[str, Any],
) -> ParsedDocument:
    segments = split_pdf_extracted_text_into_segments(text) if document_format == "pdf" else (split_text_into_segments(text) if text else [])
    raw_segments = [(segment, None, None, None) for segment in segments]
    return build_parsed_document_from_raw_segments(
        raw_segments,
        parser=parser,
        document_format=document_format,
        fallback_source=semantic_source,
        warnings=warnings,
        parse_engine=parse_engine,
        trace=trace,
    )


def split_pdf_extracted_text_into_segments(text: str) -> List[str]:
    """Convert PDF extractor line text into paragraph-like segments.

    MarkItDown can emit rendered PDF lines
    separated by single newlines. Passing those lines into the generic splitter
    turns each visual line into a paragraph and the API later serializes it as
    `line\n\nline`. For PDFs, treat single newlines inside an extracted block as
    soft wraps, while preserving explicit blank-line blocks and structural lines.
    """
    normalized = normalize_parsed_document_text(text)
    if not normalized:
        return []

    segments: List[str] = []
    for block in re.split(r"\n{2,}", normalized):
        lines = [normalize_segment_line(line) for line in block.split("\n")]
        raw_lines = [(line, None, None, None) for line in lines if line]
        if not raw_lines:
            continue
        segments.extend(segment_text for segment_text, _decision, _page, _bbox in merge_pdf_line_segments(raw_lines))
    return segments


def build_parsed_document_from_raw_segments(
    raw_segments: List[Tuple[str, Optional[SegmentSemanticDecision], Optional[int], Optional[str]]],
    *,
    parser: str,
    document_format: str,
    fallback_source: str,
    warnings: List[str],
    parse_engine: str,
    trace: Dict[str, Any],
) -> ParsedDocument:
    text = "\n\n".join(segment_text for segment_text, _decision, _page, _bbox in raw_segments)
    classifier = TextRuleSemanticClassifier(source=fallback_source)
    fallback_decisions = classifier.classify_segments([segment_text for segment_text, _decision, _page, _bbox in raw_segments])
    parsed_segments: List[ParsedSegment] = []
    cursor = 0
    for index, ((segment_text, explicit_decision, page_number, bbox_json), fallback_decision) in enumerate(zip(raw_segments, fallback_decisions)):
        decision = merge_semantic_decisions(explicit_decision, fallback_decision)
        start = cursor
        end = start + len(segment_text)
        parsed_segments.append(
            ParsedSegment(
                index=index,
                text=segment_text,
                semantic_type=decision.semantic_type,
                semantic_source=decision.semantic_source,
                semantic_confidence=decision.semantic_confidence,
                reduce_allowed=decision.reduce_allowed,
                semantic_reason=decision.semantic_reason,
                char_start=start,
                char_end=end,
                page_number=page_number,
                bbox_json=bbox_json,
            )
        )
        cursor = end + 2
    return ParsedDocument(
        text=text,
        segments=parsed_segments,
        parser=parser,
        warnings=warnings,
        document_format=document_format,
        parse_engine=parse_engine,
        parse_fallback_used=False,
        parse_trace={**trace, "segment_count": len(parsed_segments)},
    )


def merge_semantic_decisions(
    explicit_decision: Optional[SegmentSemanticDecision],
    fallback_decision: SegmentSemanticDecision,
) -> SegmentSemanticDecision:
    if (
        explicit_decision
        and explicit_decision.semantic_source == SEMANTIC_SOURCE_MINERU
        and explicit_decision.semantic_type != SEMANTIC_TYPE_SECTION_HEADING
    ):
        return explicit_decision
    # Text rules for front/back matter are more specific than generic DOCX/PDF
    # heading labels. This prevents a styled "摘要" heading from losing its
    # ABSTRACT section state and protects the following abstract body.
    if fallback_decision.semantic_type in {
        SEMANTIC_TYPE_ABSTRACT_HEADING,
        SEMANTIC_TYPE_ABSTRACT_BODY,
        SEMANTIC_TYPE_ACK_HEADING,
        SEMANTIC_TYPE_ACK_BODY,
        SEMANTIC_TYPE_REFERENCE_HEADING,
        SEMANTIC_TYPE_REFERENCE_ITEM,
        SEMANTIC_TYPE_TOC_HEADING,
        SEMANTIC_TYPE_TOC_ITEM,
        SEMANTIC_TYPE_KEYWORDS,
        SEMANTIC_TYPE_FORMULA,
        SEMANTIC_TYPE_META,
        SEMANTIC_TYPE_CAPTION,
        SEMANTIC_TYPE_MIXED_HEADING_BODY,
    }:
        return fallback_decision
    return explicit_decision or fallback_decision


def semantic_decision_from_segment(segment: Any, *, fallback_source: str = SEMANTIC_SOURCE_LEGACY_TEXT_RULE) -> SegmentSemanticDecision:
    semantic_type = getattr(segment, "semantic_type", None)
    if semantic_type:
        confidence = getattr(segment, "semantic_confidence", None)
        try:
            confidence_value = float(confidence) if confidence is not None else 0.75
        except (TypeError, ValueError):
            confidence_value = 0.75
        return SegmentSemanticDecision(
            semantic_type=str(semantic_type),
            semantic_source=str(getattr(segment, "semantic_source", None) or fallback_source),
            semantic_confidence=confidence_value,
            reduce_allowed=bool(getattr(segment, "reduce_allowed", semantic_type in REDUCE_ALLOWED_TYPES)),
            semantic_reason=str(getattr(segment, "semantic_reason", None) or "stored_semantic_type"),
            length=count_text_length(getattr(segment, "original_text", "") or ""),
        )
    return TextRuleSemanticClassifier(source=fallback_source).classify_text(getattr(segment, "original_text", "") or "")


def serialize_parse_trace(trace: Dict[str, Any]) -> str:
    return json.dumps(trace or {}, ensure_ascii=False)
